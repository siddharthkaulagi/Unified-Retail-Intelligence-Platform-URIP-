"""
Final College Project Report Generator - Phase 2
ML-Driven Unified Retail Intelligence Platform (URIP)
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

def create_urip_report():
    doc = Document()
    
    # Set margins: Left 1.25", Right 1", Top/Bottom 0.75"
    for section in doc.sections:
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
    
    def format_para(p, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.alignment = align
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(size)
            run.font.bold = bold
    
    def add_heading(text, level=1):
        if level == 1:
            p = doc.add_heading(text, level=1)
            format_para(p, 16, True, WD_ALIGN_PARAGRAPH.LEFT)
        elif level == 2:
            p = doc.add_heading(text, level=2)
            format_para(p, 16, False, WD_ALIGN_PARAGRAPH.LEFT)
        else:
            p = doc.add_heading(text, level=3)
            format_para(p, 14, False, WD_ALIGN_PARAGRAPH.LEFT)
        return p
    
    def add_para(text):
        p = doc.add_paragraph(text)
        format_para(p)
        return p
    
    # OUTER TITLE PAGE
    p = doc.add_paragraph('ML-DRIVEN UNIFIED RETAIL INTELLIGENCE PLATFORM (URIP)')
    format_para(p, 18, True, WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(8): doc.add_paragraph()
    p = doc.add_paragraph('A Project Report')
    format_para(p, 14, False, WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph('Submitted in partial fulfillment of the requirements\nfor the award of the degree of')
    format_para(p, 12, False, WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph('BACHELOR OF ENGINEERING')
    format_para(p, 14, True, WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph('in')
    format_para(p, 12, False, WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph('COMPUTER SCIENCE AND ENGINEERING')
    format_para(p, 14, True, WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph('[College Name]\n[2025]')
    format_para(p, 14, True, WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()
    
    # INNER TITLE PAGE  
    p = doc.add_paragraph('ML-DRIVEN UNIFIED RETAIL INTELLIGENCE PLATFORM (URIP)')
    format_para(p, 18, True, WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph('Submitted by\n[Student Name]\n[USN]')
    format_para(p, 14, True, WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(3): doc.add_paragraph()
    p = doc.add_paragraph('Under the guidance of\n[Guide Name]\n[Designation]')
    format_para(p, 14, True, WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(4): doc.add_paragraph()
    p = doc.add_paragraph('Department of Computer Science and Engineering\n[College Name]\n[2025]')
    format_para(p, 14, True, WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()
    
    # CERTIFICATE
    p = doc.add_paragraph('CERTIFICATE')
    format_para(p, 18, True, WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para('This is to certify that the project work entitled "ML-DRIVEN UNIFIED RETAIL INTELLIGENCE PLATFORM (URIP)" is a bonafide work carried out by [Student Name] (USN: [USN]) in partial fulfillment for the award of Bachelor of Engineering in Computer Science and Engineering of [University Name] during the academic year [Year].')
    for _ in range(4): doc.add_paragraph()
    p = doc.add_paragraph('Internal Guide\n[Guide Name]\n[Designation]')
    format_para(p, 12, True, WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph()
    p = doc.add_paragraph('Head of Department\n[HOD Name]\n[Department]')
    format_para(p, 12, True, WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()
    
    # ACKNOWLEDGMENT
    p = doc.add_paragraph('ACKNOWLEDGMENT')
    format_para(p, 18, True, WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para('I express sincere gratitude to my project guide [Guide Name] for invaluable guidance and support throughout this work. My thanks to the Head of Department [HOD Name] and faculty members of the Computer Science and Engineering department for their encouragement. I am grateful to my family and friends for their constant motivation.')
    doc.add_page_break()
    
    # ABSTRACT
    p = doc.add_paragraph('ABSTRACT')
    format_para(p, 18, True, WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para('The Unified Retail Intelligence Platform (URIP) integrates machine learning for sales forecasting, inventory optimization, and strategic planning. It employs ARIMA, Prophet, XGBoost, LightGBM, and Random Forest models with ABC/XYZ/FSN classification, GIS analysis, facility planning, and CRM analytics. Built with Streamlit, it provides interactive dashboards, automated reports, and an AI chatbot for comprehensive business intelligence.')
    doc.add_page_break()
    
    # TABLE OF CONTENTS
    p = doc.add_paragraph('TABLE OF CONTENTS')
    format_para(p, 18, True, WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    toc = [('CERTIFICATE', 'i'), ('ACKNOWLEDGMENT', 'ii'), ('ABSTRACT', 'iii'), ('LIST OF FIGURES', 'v'), ('LIST OF TABLES', 'vi'),
           ('CHAPTER 1: INTRODUCTION', '1'), ('CHAPTER 2: LITERATURE SURVEY', '7'), ('CHAPTER 3: METHODOLOGY', '17'),
           ('CHAPTER 4: SYSTEM DESIGN', '25'), ('CHAPTER 5: MODEL OVERVIEW', '37'), ('CHAPTER 6: IMPLEMENTATION', '50'),
           ('CHAPTER 7: RESULTS & DISCUSSION', '55'), ('CHAPTER 8: CONCLUSION', '60'), ('REFERENCES', '62')]
    for entry, page in toc:
        p = doc.add_paragraph(f'{entry} {"." * (70 - len(entry))} {page}')
        format_para(p, 12, False, WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()
    
    # LIST OF FIGURES
    p = doc.add_paragraph('LIST OF FIGURES')
    format_para(p, 18, True, WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    figs = [('Fig 1.1: URIP Overview', '6'), ('Fig 3.1: Preprocessing Workflow', '18'),
            ('Fig 4.1: System Architecture', '26'), ('Fig 6.1: Dashboard', '51'), ('Fig 7.1: Results', '56')]
    for fig, page in figs:
        p = doc.add_paragraph(f'{fig} {"." * (70 - len(fig))} {page}')
        format_para(p, 12, False, WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()
    
    # LIST OF TABLES
    p = doc.add_paragraph('LIST OF TABLES')
    format_para(p, 18, True, WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    tables = [('Table 5.1: Model Comparison', '49'), ('Table 7.1: Performance Metrics', '57')]
    for table, page in tables:
        p = doc.add_paragraph(f'{table} {"." * (70 - len(table))} {page}')
        format_para(p, 12, False, WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()
    
    # Save with the requested name
    output_path = 'Final_Report_Phase_2.docx'
    doc.save(output_path)
    return output_path

if __name__ == '__main__':
    result = create_urip_report()
    print(f'Report generated successfully: {result}')
