"""
Gemini AI Integration for Retail Sales Forecasting Platform
Provides AI-powered analysis and recommendations for forecast results
"""

import os
import json
import logging
from typing import Dict, Any
from datetime import datetime
from docx import Document
from dotenv import load_dotenv

# Try to import Google Generative AI
try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False
    genai = None

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------
# GEMINI AI ASSISTANT CLASS
# ---------------------------------------------------------------------

class GeminiAIAssistant:
    """Gemini AI assistant for retail forecasting analysis"""

    def __init__(self):
        """Initialize Gemini AI assistant"""
        if not GEMINI_AVAILABLE:
            raise ImportError("Google Generative AI package not installed. Please install with: pip install google-generativeai")

        api_key = os.getenv('GEMINI_API_KEY')
        if not api_key:
            raise ValueError("GEMINI_API_KEY not found in environment variables")

        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-2.5-pro')

    # -------------------------------------------------------------
    # DOCUMENT READING
    # -------------------------------------------------------------
    def read_word_document(self, file_path: str) -> str:
        """Read content from a Word document"""
        try:
            doc = Document(file_path)
            content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text.strip())
            return '\n'.join(content)
        except Exception as e:
            logger.error(f"Error reading Word document: {str(e)}")
            raise

    # -------------------------------------------------------------
    # AI ANALYSIS
    # -------------------------------------------------------------
    def analyze_forecast_results(self, forecast_content: str) -> Dict[str, Any]:
        """Analyze forecast results and provide comprehensive recommendations"""
        try:
            prompt = f"""
            You are a senior retail analytics consultant analyzing sales forecast results.
            Provide detailed insights and recommendations in JSON format based on the following data:

            FORECAST CONTENT:
            {forecast_content}

            The JSON should include these sections:
            - demand_analysis
            - inventory_management
            - sales_optimization
            - operational_efficiency
            - risk_assessment
            - strategic_recommendations
            """

            response = self.model.generate_content(prompt)
            response_text = response.text

            # Try to extract JSON
            try:
                if '```json' in response_text:
                    json_start = response_text.find('```json') + 7
                    json_end = response_text.find('```', json_start)
                    json_content = response_text[json_start:json_end].strip()
                    analysis = json.loads(json_content)
                else:
                    analysis = json.loads(response_text)
            except Exception:
                # Fallback to structured text
                analysis = self._structure_text_response(response_text)

            return {
                'success': True,
                'analysis': analysis,
                'timestamp': datetime.now().isoformat(),
                'content_length': len(forecast_content)
            }

        except Exception as e:
            logger.error(f"Error analyzing forecast: {str(e)}")
            return {'success': False, 'error': str(e)}

    # -------------------------------------------------------------
    # STRUCTURE TEXT RESPONSE
    # -------------------------------------------------------------
    def _structure_text_response(self, response_text: str) -> Dict[str, Any]:
        """Structure text response when JSON isn't returned"""
        sections = {
            'demand_analysis': '',
            'inventory_management': '',
            'sales_optimization': '',
            'operational_efficiency': '',
            'risk_assessment': '',
            'strategic_recommendations': ''
        }

        lines = response_text.split('\n')
        current_section = None

        for line in lines:
            line_lower = line.lower().strip()
            if 'demand' in line_lower:
                current_section = 'demand_analysis'
            elif 'inventory' in line_lower:
                current_section = 'inventory_management'
            elif 'sales' in line_lower:
                current_section = 'sales_optimization'
            elif 'operational' in line_lower:
                current_section = 'operational_efficiency'
            elif 'risk' in line_lower:
                current_section = 'risk_assessment'
            elif 'strategic' in line_lower:
                current_section = 'strategic_recommendations'
            elif current_section:
                sections[current_section] += line.strip() + ' '

        return sections

    # -------------------------------------------------------------
    # REPORT GENERATION
    # -------------------------------------------------------------
    def generate_enhanced_report_content(self, original_content: str, analysis: Dict) -> str:
        """Combine forecast data and AI analysis into a full enhanced report"""
        enhanced_content = f"""
# AI-ENHANCED RETAIL FORECAST ANALYSIS REPORT

## EXECUTIVE SUMMARY
This report merges forecast data with AI-driven insights for actionable retail decision-making.

---

## ORIGINAL FORECAST DATA
{original_content}

---

## AI ANALYSIS & RECOMMENDATIONS

### 1. DEMAND ANALYSIS
{analysis.get('demand_analysis', 'No data')}

### 2. INVENTORY MANAGEMENT
{analysis.get('inventory_management', 'No data')}

### 3. SALES OPTIMIZATION
{analysis.get('sales_optimization', 'No data')}

### 4. OPERATIONAL EFFICIENCY
{analysis.get('operational_efficiency', 'No data')}

### 5. RISK ASSESSMENT
{analysis.get('risk_assessment', 'No data')}

### 6. STRATEGIC RECOMMENDATIONS
{analysis.get('strategic_recommendations', 'No data')}

---

## IMPLEMENTATION ROADMAP
- **Immediate:** Apply inventory and pricing changes based on forecast.
- **Short-term (1-3 months):** Improve demand sensing and process efficiency.
- **Long-term (6-12 months):** Expand into high-demand regions and invest in analytics tools.

---

Report generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
Powered by Gemini AI
"""
        return enhanced_content

    # -------------------------------------------------------------
    # MAIN PROCESS FUNCTION
    # -------------------------------------------------------------
    def process_forecast_document(self, document_path: str) -> Dict[str, Any]:
        """Run complete pipeline on a document"""
        try:
            original_content = self.read_word_document(document_path)
            analysis_result = self.analyze_forecast_results(original_content)

            if not analysis_result.get('success'):
                return analysis_result

            enhanced_content = self.generate_enhanced_report_content(
                original_content,
                analysis_result['analysis']
            )

            return {
                'success': True,
                'analysis': analysis_result['analysis'],
                'enhanced_content': enhanced_content,
                'timestamp': datetime.now().isoformat()
            }

        except Exception as e:
            return {'success': False, 'error': str(e)}

# ---------------------------------------------------------------------
# GEMINI CHATBOT CLASS
# ---------------------------------------------------------------------

class GeminiChatbot:
    """Interactive chatbot using Gemini AI for real-time conversations"""

    def __init__(self):
        """Initialize Gemini chatbot"""
        if not GEMINI_AVAILABLE:
            raise ImportError("Google Generative AI package not installed. Please install with: pip install google-generativeai")

        api_key = os.getenv('GEMINI_API_KEY')
        if not api_key:
            raise ValueError("GEMINI_API_KEY not found in environment variables")

        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-2.5-pro')
        self.conversation_history = []

    def chat(self, message: str, context: str = None) -> Dict[str, Any]:
        """
        Send a message to the chatbot and get a response

        Args:
            message (str): User's message
            context (str): Optional context about retail/forecast data

        Returns:
            Dict containing response and metadata
        """
        try:
            # Build context-aware prompt
            if context:
                prompt = f"""
                You are an AI retail forecasting assistant. Use this context to inform your responses:

                CONTEXT:
                {context}

                USER MESSAGE:
                {message}

                Provide helpful, accurate responses about retail forecasting, sales predictions, and data analysis.
                If the user asks about forecasts or predictions, provide specific insights based on the context.
                IMPORTANT: Keep your response SHORT, CRISP, and CONCISE. Avoid long paragraphs.
                """
            else:
                prompt = f"""
                You are an AI retail forecasting assistant specializing in sales predictions and market analysis.
                
                USER MESSAGE:
                {message}
                
                Provide expert insights on retail forecasting, sales trends, inventory management, and market predictions.
                IMPORTANT: Keep your response SHORT, CRISP, and CONCISE. Avoid long paragraphs. Use bullet points where possible.
                Do not provide generic definitions unless asked. Focus on the specific question.
                """

            # Generate response
            response = self.model.generate_content(prompt)

            # Store conversation history
            self.conversation_history.append({
                'user_message': message,
                'bot_response': response.text,
                'timestamp': datetime.now().isoformat()
            })

            return {
                'success': True,
                'response': response.text,
                'timestamp': datetime.now().isoformat(),
                'conversation_length': len(self.conversation_history)
            }

        except Exception as e:
            logger.error(f"Error in chatbot: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'timestamp': datetime.now().isoformat()
            }

    def analyze_image(self, image_data, question: str = None) -> Dict[str, Any]:
        """
        Analyze an image using Gemini Vision AI

        Args:
            image_data: PIL Image, file path, or bytes data
            question (str): Optional specific question about the image

        Returns:
            Dict containing analysis results
        """
        try:
            # Configure vision model
            vision_model = genai.GenerativeModel('gemini-2.5-pro')

            if question:
                prompt = f"Analyze this image in the context of retail forecasting and sales data. {question}"
            else:
                prompt = """
                Analyze this image and provide insights relevant to retail forecasting, sales data, or business intelligence.
                Look for charts, graphs, data visualizations, or any business-related content.
                """

            # Generate analysis
            response = vision_model.generate_content([prompt, image_data])

            return {
                'success': True,
                'analysis': response.text,
                'timestamp': datetime.now().isoformat()
            }

        except Exception as e:
            logger.error(f"Error analyzing image: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'timestamp': datetime.now().isoformat()
            }

    def get_forecast_insights(self, forecast_data: str) -> Dict[str, Any]:
        """
        Get specific forecasting insights from data

        Args:
            forecast_data (str): Forecast data or metrics

        Returns:
            Dict containing insights and predictions
        """
        try:
            prompt = f"""
            Analyze the following forecast data and provide specific, actionable insights:

            FORECAST DATA:
            {forecast_data}

            Please provide:
            1. Key trends and patterns
            2. Specific predictions for the next period
            3. Risk factors and opportunities
            4. Recommended actions
            5. Confidence level in the predictions

            Format your response as structured insights.
            """

            response = self.model.generate_content(prompt)

            return {
                'success': True,
                'insights': response.text,
                'timestamp': datetime.now().isoformat()
            }

        except Exception as e:
            logger.error(f"Error getting forecast insights: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'timestamp': datetime.now().isoformat()
            }

    def clear_conversation(self):
        """Clear conversation history"""
        self.conversation_history = []

# ---------------------------------------------------------------------
# UTILITY FUNCTIONS
# ---------------------------------------------------------------------

def analyze_forecast_file(file_path: str) -> Dict[str, Any]:
    """Convenience function"""
    assistant = GeminiAIAssistant()
    return assistant.process_forecast_document(file_path)

def get_gemini_assistant() -> GeminiAIAssistant:
    """Get a Gemini AI assistant instance"""
    return GeminiAIAssistant()

# ---------------------------------------------------------------------
# MAIN SCRIPT ENTRY POINT
# ---------------------------------------------------------------------

if __name__ == "__main__":
    file_path = "forecast_report.docx"  # your input file
    output_file = "AI_Enhanced_Forecast_Report.docx"

    if not os.path.exists(file_path):
        print(f"❌ Input file not found: {file_path}")
        exit(1)

    print("🔍 Running Gemini AI Retail Forecast Analysis...\n")

    result = analyze_forecast_file(file_path)

    if result["success"]:
        print("✅ AI Analysis Complete!\n")
        print("🔹 DEMAND ANALYSIS:\n", result["analysis"].get("demand_analysis", "N/A"))
        print("\n🔹 INVENTORY MANAGEMENT:\n", result["analysis"].get("inventory_management", "N/A"))
        print("\n🔹 SALES OPTIMIZATION:\n", result["analysis"].get("sales_optimization", "N/A"))
        print("\n🔹 RISK ASSESSMENT:\n", result["analysis"].get("risk_assessment", "N/A"))
        print("\n💾 Saving enhanced report...")

        # Save enhanced report to Word file
        doc = Document()
        doc.add_heading("AI-Enhanced Retail Forecast Analysis", 0)
        doc.add_paragraph(result["enhanced_content"])
        doc.save(output_file)

        print(f"📄 Report saved as: {output_file}")

    else:
        print("❌ Error:", result.get("error"))
