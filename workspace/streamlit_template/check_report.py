from docx import Document

doc = Document(r'dev_archive\documentation\BE_Project_Report_URIP_Complete.docx')

print("=" * 60)
print("REPORT STRUCTURE ANALYSIS")
print("=" * 60)

# Get all paragraphs
all_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

# Check for required sections
required = {
    'CERTIFICATE': False,
    'ACKNOWLEDGMENT/ACKNOWLEDGEMENT': False,
    'ABSTRACT': False,
    'TABLE OF CONTENTS': False,
    'LIST OF FIGURES': False,
    'LIST OF TABLES': False,
    'CHAPTER 1': False,
    'CHAPTER 2 (LITERATURE SURVEY)': False,
    'CHAPTER 3 (METHODOLOGY)': False,
    'CHAPTER 4': False,
    'CHAPTER 5 (MODEL OVERVIEW)': False,
    'CHAPTER 6': False,
    'CHAPTER 7': False,
    'CHAPTER 8 (CONCLUSION)': False,
    'REFERENCES': False
}

# Check each section
for text in all_text[:300]:
    upper_text = text.upper()
    if 'CERTIFICATE' in upper_text:
        required['CERTIFICATE'] = True
    if 'ACKNOWLEDGMENT' in upper_text or 'ACKNOWLEDGEMENT' in upper_text:
        required['ACKNOWLEDGMENT/ACKNOWLEDGEMENT'] = True
    if upper_text == 'ABSTRACT' or upper_text.startswith('ABSTRACT'):
        required['ABSTRACT'] = True
    if 'TABLE OF CONTENTS' in upper_text or upper_text == 'CONTENTS':
        required['TABLE OF CONTENTS'] = True
    if 'LIST OF FIGURES' in upper_text:
        required['LIST OF FIGURES'] = True
    if 'LIST OF TABLES' in upper_text:
        required['LIST OF TABLES'] = True
    if 'CHAPTER 1' in upper_text:
        required['CHAPTER 1'] = True
    if 'CHAPTER 2' in upper_text:
        if 'LITERATURE' in upper_text or 'SURVEY' in upper_text:
            required['CHAPTER 2 (LITERATURE SURVEY)'] = True
    if 'CHAPTER 3' in upper_text:
        if 'METHODOLOGY' in upper_text:
            required['CHAPTER 3 (METHODOLOGY)'] = True
    if 'CHAPTER 4' in upper_text:
        required['CHAPTER 4'] = True
    if 'CHAPTER 5' in upper_text:
        if 'MODEL' in upper_text:
            required['CHAPTER 5 (MODEL OVERVIEW)'] = True
    if 'CHAPTER 6' in upper_text:
        required['CHAPTER 6'] = True
    if 'CHAPTER 7' in upper_text:
        required['CHAPTER 7'] = True
    if 'CHAPTER 8' in upper_text:
        if 'CONCLUSION' in upper_text:
            required['CHAPTER 8 (CONCLUSION)'] = True
    if 'REFERENCES' in upper_text and len(text) < 30:
        required['REFERENCES'] = True

print("\nREQUIRED SECTIONS CHECK:")
print("-" * 60)
for section, found in required.items():
    status = "✓ FOUND" if found else "✗ MISSING"
    print(f"{section:45} {status}")

print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
print(f"Total non-empty paragraphs: {len(all_text)}")

# Find all chapters
print("\n" + "=" * 60)
print("CHAPTERS FOUND:")
print("=" * 60)
chapters = [t for t in all_text if 'CHAPTER' in t.upper() and len(t) < 100]
for ch in chapters[:10]:
    print(f"  • {ch}")

# Count references
ref_section_started = False
ref_count = 0
for text in all_text:
    if 'REFERENCES' in text.upper() and len(text) < 30:
        ref_section_started = True
    elif ref_section_started:
        if text.startswith('[') and ']' in text:
            ref_count += 1
        if ref_count > 30 or 'APPENDIX' in text.upper():
            break

print(f"\nEstimated reference count: ~{ref_count}")

print("\n" + "=" * 60)
print("RECOMMENDATION:")
print("=" * 60)

missing_count = sum(1 for found in required.values() if not found)
if missing_count == 0:
    print("✓ This report appears to have ALL required sections!")
    print("  However, you should verify:")
    print("  - Formatting (Times New Roman, 1.5 spacing, margins)")
    print("  - Literature survey has 10 detailed paragraphs")
    print("  - Model descriptions are 10-20 lines each")
    print("  - 30 IEEE references are present")
elif missing_count <= 3:
    print(f"⚠ This report is mostly complete but missing {missing_count} sections.")
    print("  You may need to add these missing parts.")
else:
    print(f"✗ This report is missing {missing_count} major sections.")
    print("  Recommend generating a new report from scratch.")
