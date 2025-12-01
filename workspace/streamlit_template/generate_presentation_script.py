"""
Generate Presentation Script for URIP
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

doc = Document()

# Set margins
for section in doc.sections:
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

def format_para(p, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.alignment = align
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color

def add_speaker(name):
    p = doc.add_paragraph(name.upper())
    format_para(p, 14, True, WD_ALIGN_PARAGRAPH.LEFT, RGBColor(0, 51, 102)) # Dark Blue

def add_dialogue(text):
    p = doc.add_paragraph(text)
    format_para(p, 12, False)

def add_note(text):
    p = doc.add_paragraph(f"[{text}]")
    format_para(p, 11, True, WD_ALIGN_PARAGRAPH.LEFT, RGBColor(100, 100, 100)) # Grey

# TITLE
p = doc.add_paragraph('FINAL YEAR PROJECT PRESENTATION SCRIPT')
format_para(p, 18, True, WD_ALIGN_PARAGRAPH.CENTER)
p = doc.add_paragraph('ML-Driven Unified Retail Intelligence Platform (URIP)')
format_para(p, 16, False, WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

add_note('Total Time: 10-12 Minutes')
add_note('Speakers: Siddharth, Charan, Lohit')
doc.add_paragraph()
doc.add_page_break()

# === SIDDHARTH ===
add_speaker('SIDDHARTH (Introduction & System Overview)')
add_note('Time: ~3-4 Minutes')

add_dialogue("Good morning to the respected panel and my dear friends. We are here to present our final year project titled 'ML-Driven Unified Retail Intelligence Platform', or URIP for short. This is a web-based decision support system designed specifically for small and growing retailers.")

add_dialogue("Let me start with the problem statement. In today's market, large retail giants like Reliance or DMart have massive data science teams to optimize their stock, predict sales, and choose the best store locations. However, small and medium retailers—like your local supermarket or a growing chain like KPN Fresh—often rely on intuition or basic Excel sheets. They face three critical issues: First, they struggle to forecast demand accurately, leading to either stockouts or wasted inventory. Second, their inventory management is reactive rather than proactive. And third, when they want to expand, they lack the data to choose the right location.")

add_dialogue("Our objective was to bridge this gap. We wanted to build a single, unified platform that democratizes advanced analytics. Our goal was to integrate Sales Forecasting, Inventory Classification, GIS Location Intelligence, and CRM into one easy-to-use web application.")

add_dialogue("Let me give you a high-level overview of the system architecture. The workflow begins with Data Ingestion, where the retailer uploads their transaction history. This data goes through a Preprocessing Pipeline to handle missing values and outliers. It is then fed into our core analytical engines: the Forecasting Engine, the Inventory Classifier, and the GIS Module. Finally, all these insights are visualized on an interactive Dashboard built using Streamlit, which allows the user to download comprehensive reports.")

add_dialogue("So, URIP isn't just a model running in a notebook; it is a fully functional, end-to-end web application that transforms raw data into actionable business decisions.")

add_dialogue("Now, I will hand over to Charan to explain the methodology and the machine learning models we implemented.")

doc.add_paragraph()
doc.add_paragraph()

# === CHARAN ===
add_speaker('CHARAN (Methodology, Models & GIS Logic)')
add_note('Time: ~3-4 Minutes')

add_dialogue("Thank you, Siddharth. Good morning everyone. I will walk you through the technical core of our project.")

add_dialogue("For the Sales Forecasting module, we didn't rely on just one algorithm. We implemented a multi-model approach because retail data is complex. We used ARIMA to capture linear trends and Prophet to handle strong seasonal patterns and holidays. To capture non-linear relationships and complex interactions, we implemented tree-based models like XGBoost, LightGBM, and Random Forest. Finally, we used a Weighted Ensemble technique, which combines the predictions of all these models to give us the highest possible accuracy.")

add_dialogue("To evaluate these models, we used metrics like RMSE (Root Mean Squared Error) and MAPE (Mean Absolute Percentage Error). This ensures that the forecasts we present to the user are statistically validated.")

add_dialogue("Moving to Inventory Analytics, we implemented a multi-dimensional classification system. We don't just look at sales volume. We perform ABC Analysis to identify high-value items—so the retailer knows which 20% of items give 80% of revenue. We use XYZ Analysis to measure demand variability—telling them which items are predictable versus volatile. And we use FSN Analysis to classify items as Fast, Slow, or Non-moving. By combining these, we generate a strategic matrix that tells the store manager exactly how to handle each product category.")

add_dialogue("For the GIS Module, we used GeoPandas and Folium. The logic here is spatial analysis. We map the retailer's existing stores and competitor locations using latitude and longitude. We then perform Buffer Analysis to visualize coverage areas. We also integrated ward-wise population data of Bangalore. By combining population density with competitor proximity, our algorithm calculates a 'Location Score' to recommend the best zones for opening new stores.")

add_dialogue("I will now request Lohit to demonstrate the actual application and discuss the results.")

doc.add_paragraph()
doc.add_paragraph()

# === LOHIT ===
add_speaker('LOHIT (Demo, Results & Conclusion)')
add_note('Time: ~3-4 Minutes')

add_dialogue("Thank you, Charan. Good morning. I will now take you through the working of the URIP application and our results.")

add_dialogue("Our platform is built using Python and Streamlit, leveraging Plotly for interactive charts and Folium for maps. The user experience is designed to be simple and intuitive.")

add_dialogue("When a user logs in, they land on the Home Dashboard, which gives an immediate snapshot of key KPIs like Total Revenue, Top Selling Products, and Recent Trends.")

add_dialogue("The workflow is straightforward: The user goes to the 'Data Upload' page and uploads their sales CSV file. The system automatically processes this data. They can then navigate to the 'Forecasting' module. Here, they can select a product category, choose a model—say, Prophet or the Ensemble—and instantly see the sales prediction for the next 12 weeks. The interactive graphs show both historical data and the future forecast with confidence intervals.")

add_dialogue("Next, in the 'Inventory Analytics' tab, the user can see the ABC-XYZ classification. For example, they can filter for 'AX' items—high value, low variability—which are their cash cows. This helps them prioritize stock replenishment.")

add_dialogue("In the 'GIS Analysis' module, we visualize the store network. The map shows KPN Fresh stores in green and competitors in red. We can overlay population heatmaps to see which high-density areas are underserved, providing data-backed expansion strategies.")

add_dialogue("Finally, the 'Reports' module allows the user to generate a full project report or a business summary in Word format with a single click, which is crucial for documentation.")

add_dialogue("In conclusion, URIP successfully demonstrates that powerful machine learning tools can be made accessible to smaller retailers. Our Ensemble model achieved an average MAPE of around 8-10%, which is highly effective for retail planning. The platform empowers store owners to move away from guesswork and make data-driven decisions regarding stock, sales, and expansion.")

add_dialogue("For future scope, we plan to integrate real-time data streaming from POS systems and develop a mobile version for field managers.")

add_dialogue("That concludes our presentation. We are now open to any questions. Thank you.")

# Save document
doc.save('URIP_Presentation_Script.docx')
