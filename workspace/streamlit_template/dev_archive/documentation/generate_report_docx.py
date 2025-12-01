"""
Script to convert the 4-part BE Project Report from Markdown to Word Document
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_markdown_content(doc, markdown_file):
    """Read markdown file and add to Word document with basic formatting"""
    
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            continue
            
        # Handle headers
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('#### '):
            heading = doc.add_heading(line[5:], level=4)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Handle bullet points
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        
        # Handle numbered lists
        elif line[0:2].replace('.', '').isdigit():
            doc.add_paragraph(line.split('. ', 1)[1] if '. ' in line else line, 
                            style='List Number')
        
        # Handle code blocks (simplified)
        elif line.startswith('```'):
            continue
        
        # Regular paragraph
        else:
            if line.startswith('**') or line.startswith('*'):
                # Contains bold/italic - add as normal paragraph
                p = doc.add_paragraph(line)
            else:
                p = doc.add_paragraph(line)

def create_word_report():
    """Create complete Word document from 4 markdown parts"""
    
    # Initialize document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Add title page
    title = doc.add_heading('B.E. PROJECT REPORT', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    main_title = doc.add_heading('UNIFIED RETAIL INTELLIGENCE PLATFORM (URIP)', level=1)
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('ML-Driven Retail Supply Chain Optimization and Sales Forecasting System')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # File paths
    artifact_dir = r'C:\Users\sidda\.gemini\antigravity\brain\9a49d5ea-491d-464c-9fa2-ad33dc96ed4c'
    
    parts = [
        os.path.join(artifact_dir, 'BE_Project_Report_Part1.md'),
        os.path.join(artifact_dir, 'BE_Project_Report_Part2.md'),
        os.path.join(artifact_dir, 'BE_Project_Report_Part3.md'),
        os.path.join(artifact_dir, 'BE_Project_Report_Part4.md')
    ]
    
    # Add content from each part
    for i, part_file in enumerate(parts, 1):
        print(f"Processing Part {i}...")
        if os.path.exists(part_file):
            add_markdown_content(doc, part_file)
            if i < 4:  # Don't add page break after last part
                doc.add_page_break()
        else:
            print(f"Warning: {part_file} not found!")
    
    # Save document
    output_path = r'C:\Users\sidda\Downloads\Retail Sales Prediction\workspace\streamlit_template\BE_Project_Report_URIP_Complete.docx'
    doc.save(output_path)
    print(f"\n✅ Report saved successfully!")
    print(f"📄 Location: {output_path}")
    print(f"\nYou can now open and edit this Word document.")
    
    return output_path

if __name__ == "__main__":
    try:
        output_file = create_word_report()
        print("\n" + "="*60)
        print("SUCCESS! Your BE Project Report is ready.")
        print("="*60)
    except Exception as e:
        print(f"\n❌ Error: {e}")
        print("\nTroubleshooting:")
        print("1. Make sure python-docx is installed: pip install python-docx")
        print("2. Check that all 4 markdown files exist in the artifacts folder")
