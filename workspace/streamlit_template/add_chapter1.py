"""
Add all chapters to Final Report Phase 2
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

doc = Document('Final_Report_Phase_2.docx')

def format_para(p, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.alignment = align
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.bold = bold

def add_heading(text, level=1):
    if level == 1:
        p = doc.add_heading(text, level=1)
        format_para(p, 16, True, WD_ALIGN_PARAGRAPH.LEFT)
    elif level == 2:
        p = doc.add_heading(text, level=2)
        format_para(p, 16, False, WD_ALIGN_PARAGRAPH.LEFT)
    else:
        p = doc.add_heading(text, level=3)
        format_para(p, 14, False, WD_ALIGN_PARAGRAPH.LEFT)
    return p

def add_para(text):
    p = doc.add_paragraph(text)
    format_para(p)
    return p

# CHAPTER 1: INTRODUCTION
add_heading('CHAPTER 1: INTRODUCTION', 1)
doc.add_paragraph()

add_heading('1.1 Background', 2)
add_para('The retail industry faces complex challenges in inventory management, demand forecasting, and strategic planning. Traditional methods prove inadequate for modern data-driven environments. The exponential growth of transactional data and advances in machine learning create opportunities for actionable insights. However, many businesses lack integrated platforms to harness ML power. The Unified Retail Intelligence Platform bridges this gap by combining forecasting models, GIS analytics, inventory optimization, and CRM into a cohesive system, empowering decision-makers with cutting-edge analytical capabilities accessible through an intuitive interface.')

add_heading('1.2 Motivation', 2)
add_para('Sales forecasting remains critical yet challenging—inaccurate predictions cause stockouts or excess inventory. Inventory management complexities vary across categories and locations. Strategic location planning for expansions requires sophisticated demographic and competitor analysis typically available only to large enterprises. Customer data proliferation creates CRM opportunities that remain underutilized. Traditional BI tools operate in silos, requiring multiple systems for comprehensive insights. High technical barriers limit ML adoption. These challenges motivated creating a unified platform democratizing advanced retail analytics for businesses of all sizes.')

add_heading('1.3 Problem Description', 2)
add_para('Retail businesses face demand uncertainty with complex sales patterns including seasonality, promotions, and market shifts that defeat simple statistical methods. Inventory optimization requires balancing holding costs against availability. Geographic expansion decisions often rely on intuition rather than data analysis. Facility layouts evolve organically causing inefficiencies. Customer segmentation uses simple demographics instead of behavioral analytics. Managers lack real-time KPI visibility. No integrated system addresses these problems holistically, forcing businesses to manually integrate insights across disparate analytical workflows.')

add_heading('1.4 Objectives', 2)
add_para('Primary objectives include: (1) Develop integrated ML platform combining forecasting, inventory analytics, GIS, and CRM. (2) Implement and compare ARIMA, Prophet, XGBoost, LightGBM, Random Forest with automatic model selection. (3) Create ABC/XYZ/FSN inventory classification system. (4) Integrate GIS for store location analysis and expansion opportunities. (5) Provide facility layout planning tools. (6) Implement RFM customer segmentation. (7) Develop intuitive dashboard interface. (8) Enable automated report generation. (9) Integrate AI-powered chatbot. (10) Ensure platform scalability and extensibility.')

add_heading('1.5 Scope', 2)
add_para('URIP uses Python with scikit-learn, Prophet, XGBoost, LightGBM, GeoPandas, Folium, and Plotly. Streamlit provides the UI framework. Supports CSV/Excel data ingestion with preprocessing pipelines. Forecast capabilities extend to customizable horizons with confidence intervals. GIS module handles KML/CSV formats for population analysis and competitive mapping. Inventory analytics covers three classification methodologies. CRM implements clustering and RFM segmentation. Targets retail businesses across grocery, supermarket, specialty retail, and multi-location chains. Designed for store managers, procurement officers, operations managers, marketing teams, and executives. Handles datasets from thousands to millions of records with scalability for business expansion.')

add_heading('1.6 Applications', 2)  
add_para('Applications span the retail value chain. Demand forecasting optimizes procurement and reduces stockouts. ABC-XYZ-FSN classification improves working capital efficiency. GIS identifies underserved markets for expansion. Layout optimization minimizes travel distances and improves space utilization. CRM enables targeted marketing and customer retention. Automated reporting supports stakeholder meetings and compliance. AI chatbot provides instant insights. Educational institutions use URIP for business analytics courses. SMEs benefit from enterprise-grade analytics without expensive software or dedicated data science teams.')
add_para('<Insert Fig 1.1: URIP System Overview Diagram>')
doc.add_page_break()

# Save progress
doc.save('Final_Report_Phase_2.docx')
print("Chapter 1 added successfully!")
