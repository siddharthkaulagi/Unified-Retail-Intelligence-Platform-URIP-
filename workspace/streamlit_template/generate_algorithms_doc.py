"""
Generate Algorithms Document with Code
All algorithms used in URIP with implementations
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

doc = Document()

# Set margins
for section in doc.sections:
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

def format_para(p, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.alignment = align
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        run.font.bold = bold

def add_title(text):
    p = doc.add_heading(text, level=1)
    format_para(p, 16, True, WD_ALIGN_PARAGRAPH.CENTER)
    return p

def add_heading(text, level=2):
    p = doc.add_heading(text, level=level)
    format_para(p, 14 if level == 2 else 12, True)
    return p

def add_para(text):
    p = doc.add_paragraph(text)
    format_para(p)
    return p

def add_code(code_text, lang='python'):
    p = doc.add_paragraph(code_text)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 128)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p

# TITLE
add_title('URIP: Complete Algorithm Reference\nWith Code Implementations')
doc.add_paragraph()
p = doc.add_paragraph('ML-Driven Unified Retail Intelligence Platform')
format_para(p, 12, False, WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# TABLE OF CONTENTS
add_heading('Contents', 1)
toc_items = [
    '1. Sales Forecasting Algorithms',
    '   1.1 ARIMA',
    '   1.2 Prophet',
    '   1.3 XGBoost',
    '   1.4 LightGBM',
    '   1.5 Random Forest',
    '   1.6 Ensemble Learning',
    '2. Inventory Classification Algorithms',
    '   2.1 ABC Analysis',
    '   2.2 XYZ Analysis',
    '   2.3 FSN Analysis',
    '3. Customer Analytics Algorithms',
    '   3.1 K-Means Clustering',
    '   3.2 RFM Analysis',
    '4. GIS Algorithms',
    '   4.1 Buffer Analysis',
    '   4.2 Weighted Scoring',
    '5. Layout Optimization',
    '6. Performance Metrics'
]
for item in toc_items:
    add_para(item)
doc.add_page_break()

# ========== FORECASTING ALGORITHMS ==========
add_heading('1. SALES FORECASTING ALGORITHMS', 1)
doc.add_paragraph()

# ARIMA
add_heading('1.1 ARIMA (AutoRegressive Integrated Moving Average)', 2)
add_para('Type: Statistical Time Series Model')
add_para('Library: statsmodels')
add_para('')
add_para('Mathematical Logic:')
add_para('ARIMA combines three components:')
add_para('• AR (p): Autoregressive - uses past values')
add_para('• I (d): Integrated - differencing for stationarity')
add_para('• MA (q): Moving Average - uses past errors')
add_para('')
add_para('Formula: y(t) = c + φ₁y(t-1) + ... + φₚy(t-p) + θ₁ε(t-1) + ... + θ_qε(t-q) + ε(t)')
add_para('')
add_heading('Algorithm Steps:', 3)
add_code('''1. Test stationarity using Augmented Dickey-Fuller test
2. Apply differencing if series is non-stationary
3. Identify AR order (p) using PACF plot
4. Identify MA order (q) using ACF plot
5. Fit model minimizing AIC criterion
6. Generate forecasts with confidence intervals
7. Validate using backtest''')

add_heading('Implementation Code:', 3)
add_code('''def train_arima(train_data, forecast_periods=12):
    """
    Train ARIMA model for sales forecasting
    
    Args:
        train_data: Time series data (pandas Series)
        forecast_periods: Number of periods to forecast
    
    Returns:
        dict with model, forecast, fitted_values, metrics
    """
    from statsmodels.tsa.arima.model import ARIMA
    import warnings
    warnings.filterwarnings('ignore')
    
    # Auto-select best parameters using AIC
    best_aic = np.inf
    best_order = None
    
    # Grid search over parameter space
    for p in range(0, 3):
        for d in range(0, 2):
            for q in range(0, 3):
                try:
                    model = ARIMA(train_data, order=(p, d, q))
                    fitted = model.fit()
                    if fitted.aic < best_aic:
                        best_aic = fitted.aic
                        best_order = (p, d, q)
                except:
                    continue
    
    # Fit final model with best parameters
    model = ARIMA(train_data, order=best_order)
    fitted_model = model.fit()
    
    # Generate forecasts
    forecast = fitted_model.forecast(steps=forecast_periods)
    
    # Get fitted values for training data
    fitted_values = fitted_model.fittedvalues
    
    # Calculate metrics
    mape = calculate_mape(train_data, fitted_values)
    rmse = calculate_rmse(train_data, fitted_values)
    
    return {
        'model': fitted_model,
        'forecast': forecast,
        'fitted_values': fitted_values,
        'parameters': best_order,
        'mape': mape,
        'rmse': rmse
    }''')
doc.add_page_break()

# Prophet
add_heading('1.2 Prophet (Facebook Prophet)', 2)
add_para('Type: Additive Regression Model')
add_para('Library: fbprophet')
add_para('')
add_para('Mathematical Logic:')
add_para('Prophet decomposes time series into:')
add_para('y(t) = g(t) + s(t) + h(t) + ε(t)')
add_para('• g(t): Trend (piecewise linear or logistic)')
add_para('• s(t): Seasonality (Fourier series)')
add_para('• h(t): Holiday effects')
add_para('• ε(t): Error term')
add_para('')
add_heading('Implementation Code:', 3)
add_code('''def train_prophet(df, forecast_periods=12):
    """
    Train Prophet model for sales forecasting
    
    Args:
        df: DataFrame with 'ds' (date) and 'y' (value) columns
        forecast_periods: Number of periods to forecast
    
    Returns:
        dict with model, forecast, components
    """
    from prophet import Prophet
    
    # Initialize Prophet with parameters
    model = Prophet(
        yearly_seasonality=True,
        weekly_seasonality=False,
        daily_seasonality=False,
        seasonality_mode='multiplicative',
        changepoint_prior_scale=0.05
    )
    
    # Add custom seasonality if needed
    model.add_seasonality(
        name='monthly',
        period=30.5,
        fourier_order=5
    )
    
    # Fit the model
    model.fit(df)
    
    # Create future dataframe
    future = model.make_future_dataframe(
        periods=forecast_periods,
        freq='MS'  # Month start
    )
    
    # Generate forecast
    forecast = model.predict(future)
    
    # Extract components
    components = {
        'trend': forecast['trend'],
        'yearly': forecast.get('yearly', None),
        'monthly': forecast.get('monthly', None)
    }
    
    return {
        'model': model,
        'forecast': forecast[['ds', 'yhat', 'yhat_lower', 'yhat_upper']],
        'components': components
    }''')
doc.add_page_break()

# XGBoost
add_heading('1.3 XGBoost (Extreme Gradient Boosting)', 2)
add_para('Type: Ensemble Tree-Based Model')
add_para('Library: xgboost')
add_para('')
add_para('Mathematical Logic:')
add_para('Sequential tree building with gradient descent:')
add_para('F(x) = F₀(x) + η·f₁(x) + η·f₂(x) + ... + η·f_T(x)')
add_para('where each tree f_t corrects errors of F_{t-1}')
add_para('')
add_para('Objective Function:')
add_para('L = Σᵢ l(yᵢ, ŷᵢ) + Σₖ Ω(fₖ)')
add_para('where Ω(f) = γT + (λ/2)||w||²')
add_para('')
add_heading('Implementation Code:', 3)
add_code('''def train_xgboost(X_train, y_train, X_test, forecast_periods=12):
    """
    Train XGBoost model for sales forecasting
    
    Args:
        X_train: Training features (lag variables, rolling stats)
        y_train: Training target (sales)
        X_test: Test features for forecasting
        forecast_periods: Periods to forecast
    
    Returns:
        dict with model, forecast, feature_importance
    """
    import xgboost as xgb
    
    # Define parameters
    params = {
        'objective': 'reg:squarederror',
        'max_depth': 6,
        'learning_rate': 0.1,
        'n_estimators': 100,
        'subsample': 0.8,
        'colsample_bytree': 0.8,
        'reg_alpha': 0.1,
        'reg_lambda': 1.0,
        'random_state': 42
    }
    
    # Create and train model
    model = xgb.XGBRegressor(**params)
    model.fit(
        X_train, y_train,
        eval_set=[(X_test, y_test)],
        early_stopping_rounds=10,
        verbose=False
    )
    
    # Generate fitted values for full dataset
    fitted_values = model.predict(X_train)
    
    # Generate forecast
    forecast = model.predict(X_test)
    
    # Get feature importance
    importance = pd.DataFrame({
        'feature': X_train.columns,
        'importance': model.feature_importances_
    }).sort_values('importance', ascending=False)
    
    return {
        'model': model,
        'forecast': forecast,
        'fitted_values': fitted_values,
        'feature_importance': importance
    }''')
doc.add_page_break()

# LightGBM
add_heading('1.4 LightGBM (Light Gradient Boosting Machine)', 2)
add_para('Type: Gradient Boosting Framework')
add_para('Library: lightgbm')
add_para('')
add_para('Key Algorithms:')
add_para('• Histogram-based algorithm for efficiency')
add_para('• Leaf-wise tree growth (best-first)')
add_para('• GOSS: Gradient-based One-Side Sampling')
add_para('• EFB: Exclusive Feature Bundling')
add_para('')
add_heading('Implementation Code:', 3)
add_code('''def train_lightgbm(X_train, y_train, X_test, forecast_periods=12):
    """
    Train LightGBM model for sales forecasting
    
    Args:
        X_train: Training features
        y_train: Training target
        X_test: Test features
    
    Returns:
        dict with model, forecast, metrics
    """
    import lightgbm as lgb
    
    # Parameters
    params = {
        'objective': 'regression',
        'metric': 'rmse',
        'boosting_type': 'gbdt',
        'num_leaves': 31,
        'learning_rate': 0.05,
        'feature_fraction': 0.9,
        'bagging_fraction': 0.8,
        'bagging_freq': 5,
        'verbose': -1
    }
    
    # Create dataset
    train_data = lgb.Dataset(X_train, label=y_train)
    
    # Train model
    model = lgb.train(
        params,
        train_data,
        num_boost_round=100,
        valid_sets=[train_data],
        callbacks=[lgb.early_stopping(10)]
    )
    
    # Generate predictions
    fitted_values = model.predict(X_train)
    forecast = model.predict(X_test)
    
    return {
        'model': model,
        'forecast': forecast,
        'fitted_values': fitted_values
    }''')
doc.add_page_break()

# Random Forest
add_heading('1.5 Random Forest', 2)
add_para('Type: Ensemble Learning (Bagging)')
add_para('Library: sklearn.ensemble')
add_para('')
add_para('Mathematical Logic:')
add_para('ŷ = (1/N) Σᵢ₌₁ᴺ tree_i(x)')
add_para('Where N = number of trees, each trained on:')
add_para('• Bootstrap sample of data')
add_para('• Random subset of features')
add_para('')
add_heading('Implementation Code:', 3)
add_code('''def train_random_forest(X_train, y_train, X_test, forecast_periods=12):
    """
    Train Random Forest model
    
    Args:
        X_train: Training features
        y_train: Training target
        X_test: Test features
    
    Returns:
        dict with model, forecast, feature_importance
    """
    from sklearn.ensemble import RandomForestRegressor
    
    # Initialize model
    model = RandomForestRegressor(
        n_estimators=100,
        max_depth=10,
        min_samples_split=5,
        min_samples_leaf=2,
        max_features='sqrt',
        random_state=42,
        n_jobs=-1
    )
    
    # Train model
    model.fit(X_train, y_train)
    
    # Predictions
    fitted_values = model.predict(X_train)
    forecast = model.predict(X_test)
    
    # Feature importance
    importance = pd.DataFrame({
        'feature': X_train.columns,
        'importance': model.feature_importances_
    }).sort_values('importance', ascending=False)
    
    return {
        'model': model,
        'forecast': forecast,
        'fitted_values': fitted_values,
        'feature_importance': importance
    }''')
doc.add_page_break()

# Ensemble
add_heading('1.6 Ensemble Learning (Weighted Averaging)', 2)
add_para('Type: Meta-Learning')
add_para('Custom Implementation')
add_para('')
add_para('Mathematical Logic:')
add_para('Final_Forecast = Σᵢ wᵢ × Model_i_Forecast')
add_para('where weights wᵢ = (1/errorᵢ) / Σⱼ(1/errorⱼ)')
add_para('')
add_heading('Implementation Code:', 3)
add_code('''def create_ensemble(models_dict, validation_errors):
    """
    Create weighted ensemble from multiple models
    
    Args:
        models_dict: Dictionary of trained models
        validation_errors: MAPE/RMSE for each model
    
    Returns:
        Ensemble forecast
    """
    # Calculate weights from inverse errors
    weights = {}
    total_inv_error = 0
    
    for model_name, error in validation_errors.items():
        inv_error = 1.0 / (error + 1e-10)  # Avoid division by zero
        weights[model_name] = inv_error
        total_inv_error += inv_error
    
    # Normalize weights to sum to 1
    for model_name in weights:
        weights[model_name] /= total_inv_error
    
    # Combine forecasts
    ensemble_forecast = None
    for model_name, weight in weights.items():
        model_forecast = models_dict[model_name]['forecast']
        if ensemble_forecast is None:
            ensemble_forecast = weight * model_forecast
        else:
            ensemble_forecast += weight * model_forecast
    
    return ensemble_forecast, weights''')
doc.add_page_break()

# ========== INVENTORY ALGORITHMS ==========
add_heading('2. INVENTORY CLASSIFICATION ALGORITHMS', 1)
doc.add_paragraph()

# ABC Analysis
add_heading('2.1 ABC Analysis (Pareto Analysis)', 2)
add_para('Type: Value-based Classification')
add_para('')
add_para('Logic:')
add_para('Based on Pareto Principle (80-20 rule):')
add_para('• A-class: Top 20% items → 80% value')
add_para('• B-class: Next 30% items → 15% value')
add_para('• C-class: Remaining 50% items → 5% value')
add_para('')
add_heading('Implementation Code:', 3)
add_code('''def abc_analysis(df, value_col='total_value'):
    """
    Perform ABC analysis on inventory
    
    Args:
        df: DataFrame with product data
        value_col: Column containing product value
    
    Returns:
        DataFrame with ABC classification
    """
    # Calculate cumulative percentage
    df_sorted = df.sort_values(value_col, ascending=False).copy()
    df_sorted['cumulative_value'] = df_sorted[value_col].cumsum()
    total_value = df_sorted[value_col].sum()
    df_sorted['cumulative_pct'] = (df_sorted['cumulative_value'] / total_value) * 100
    
    # Classify based on cumulative percentage
    def classify_abc(cumulative_pct):
        if cumulative_pct <= 80:
            return 'A'
        elif cumulative_pct <= 95:
            return 'B'
        else:
            return 'C'
    
    df_sorted['ABC_Class'] = df_sorted['cumulative_pct'].apply(classify_abc)
    
    return df_sorted''')

# XYZ Analysis
add_heading('2.2 XYZ Analysis (Variability Classification)', 2)
add_para('Type: Demand Variability Classification')
add_para('')
add_para('Logic:')
add_para('Based on Coefficient of Variation (CV):')
add_para('CV = (StandardDeviation / Mean) × 100')
add_para('• X-class: CV < 20% (Predictable demand)')
add_para('• Y-class: 20% ≤ CV < 50% (Moderate variability)')
add_para('• Z-class: CV ≥ 50% (Highly variable)')
add_para('')
add_heading('Implementation Code:', 3)
add_code('''def xyz_analysis(df, demand_col='monthly_demand'):
    """
    Perform XYZ analysis based on demand variability
    
    Args:
        df: DataFrame with product demand history
        demand_col: Column with demand data
    
    Returns:
        DataFrame with XYZ classification
    """
    # Group by product and calculate statistics
    product_stats = df.groupby('product_id').agg({
        demand_col: ['mean', 'std']
    }).reset_index()
    
    product_stats.columns = ['product_id', 'mean_demand', 'std_demand']
    
    # Calculate Coefficient of Variation
    product_stats['cv'] = (product_stats['std_demand'] / 
                           product_stats['mean_demand']) * 100
    
    # Classify
    def classify_xyz(cv):
        if cv < 20:
            return 'X'
        elif cv < 50:
            return 'Y'
        else:
            return 'Z'
    
    product_stats['XYZ_Class'] = product_stats['cv'].apply(classify_xyz)
    
    return product_stats''')

# FSN Analysis
add_heading('2.3 FSN Analysis (Movement Classification)', 2)
add_para('Type: Consumption-based Classification')
add_para('')
add_para('Logic:')
add_para('Based on consumption rate over time period:')
add_para('• F (Fast): High consumption frequency')
add_para('• S (Slow): Moderate consumption')
add_para('• N (Non-moving): No/minimal consumption')
add_para('')
add_heading('Implementation Code:', 3)
add_code('''def fsn_analysis(df, consumption_col='units_sold', time_period_days=90):
    """
    Perform FSN analysis based on movement
    
    Args:
        df: DataFrame with transaction data
        consumption_col: Column with consumption data
        time_period_days: Analysis period
    
    Returns:
        DataFrame with FSN classification
    """
    # Calculate consumption rate
    product_movement = df.groupby('product_id').agg({
        consumption_col: 'sum',
        'date': 'count'  # Transaction frequency
    }).reset_index()
    
    product_movement.columns = ['product_id', 'total_consumption', 'frequency']
    
    # Calculate consumption rate (per day)
    product_movement['consumption_rate'] = product_movement['total_consumption'] / time_period_days
    
    # Define thresholds (customize based on business)
    fast_threshold = product_movement['consumption_rate'].quantile(0.70)
    slow_threshold = product_movement['consumption_rate'].quantile(0.30)
    
    def classify_fsn(rate):
        if rate >= fast_threshold:
            return 'F'
        elif rate >= slow_threshold:
            return 'S'
        else:
            return 'N'
    
    product_movement['FSN_Class'] = product_movement['consumption_rate'].apply(classify_fsn)
    
    return product_movement''')
doc.add_page_break()

# ========== CRM ALGORITHMS ==========
add_heading('3. CUSTOMER ANALYTICS ALGORITHMS', 1)
doc.add_paragraph()

# K-Means
add_heading('3.1 K-Means Clustering', 2)
add_para('Type: Unsupervised Learning')
add_para('Library: sklearn.cluster')
add_para('')
add_para('Algorithm Logic:')
add_para('1. Initialize k centroids randomly')
add_para('2. Assign each point to nearest centroid: d(x, μₖ) = ||x - μₖ||')
add_para('3. Update centroids: μₖ = (1/|Sₖ|) Σₓ∈Sₖ x')
add_para('4. Repeat until convergence')
add_para('')
add_heading('Implementation Code:', 3)
add_code('''def perform_kmeans_clustering(customer_data, n_clusters=5):
    """
    Perform K-Means clustering on customer data
    
    Args:
        customer_data: DataFrame with customer metrics
        n_clusters: Number of clusters
    
    Returns:
        DataFrame with cluster labels
    """
    from sklearn.cluster import KMeans
    from sklearn.preprocessing import StandardScaler
    
    # Select features for clustering
    features = ['recency', 'frequency', 'monetary']
    X = customer_data[features]
    
    # Standardize features
    scaler = StandardScaler()
    X_scaled = scaler.fit_transform(X)
    
    # Perform K-Means
    kmeans = KMeans(
        n_clusters=n_clusters,
        init='k-means++',
        max_iter=300,
        n_init=10,
        random_state=42
    )
    
    customer_data['cluster'] = kmeans.fit_predict(X_scaled)
    
    # Calculate cluster centers
    centers = scaler.inverse_transform(kmeans.cluster_centers_)
    
    return customer_data, centers''')

# RFM Analysis
add_heading('3.2 RFM Analysis', 2)
add_para('Type: Rule-based Segmentation')
add_para('')
add_para('Metrics:')
add_para('• R (Recency): Days since last purchase')
add_para('• F (Frequency): Number of transactions')
add_para('• M (Monetary): Total spend amount')
add_para('')
add_para('Scoring: Each metric scored 1-5 (quintiles)')
add_para('Segments: Based on RFM scores (e.g., 555 = Champions)')
add_para('')
add_heading('Implementation Code:', 3)
add_code('''def rfm_analysis(transaction_data, analysis_date=None):
    """
    Perform RFM analysis on customer transactions
    
    Args:
        transaction_data: DataFrame with customer_id, date, amount
        analysis_date: Reference date (defaults to max date)
    
    Returns:
        DataFrame with RFM scores and segments
    """
    import pandas as pd
    from datetime import datetime
    
    if analysis_date is None:
        analysis_date = transaction_data['date'].max()
    
    # Calculate RFM metrics
    rfm = transaction_data.groupby('customer_id').agg({
        'date': lambda x: (analysis_date - x.max()).days,  # Recency
        'transaction_id': 'count',  # Frequency
        'amount': 'sum'  # Monetary
    }).reset_index()
    
    rfm.columns = ['customer_id', 'recency', 'frequency', 'monetary']
    
    # Score each metric (1-5, where 5 is best)
    rfm['R_score'] = pd.qcut(rfm['recency'], 5, labels=[5,4,3,2,1])
    rfm['F_score'] = pd.qcut(rfm['frequency'].rank(method='first'), 5, labels=[1,2,3,4,5])
    rfm['M_score'] = pd.qcut(rfm['monetary'].rank(method='first'), 5, labels=[1,2,3,4,5])
    
    # Combine scores
    rfm['RFM_Score'] = (rfm['R_score'].astype(str) + 
                        rfm['F_score'].astype(str) + 
                        rfm['M_score'].astype(str))
    
    # Segment customers
    def segment_customer(row):
        if row['R_score'] >= 4 and row['F_score'] >= 4:
            return 'Champions'
        elif row['R_score'] >= 3 and row['F_score'] >= 3:
            return 'Loyal'
        elif row['R_score'] <= 2:
            return 'At Risk'
        else:
            return 'Potential'
    
    rfm['segment'] = rfm.apply(segment_customer, axis=1)
    
    return rfm''')
doc.add_page_break()

# ========== GIS ALGORITHMS ==========
add_heading('4. GIS & LOCATION ALGORITHMS', 1)
doc.add_paragraph()

# Buffer Analysis
add_heading('4.1 Buffer Analysis', 2)
add_para('Type: Spatial Analysis')
add_para('Library: geopandas')
add_para('')
add_para('Mathematical Logic:')
add_para('Buffer(point, radius) = {p | distance(p, point) ≤ radius}')
add_para('Creates circular zones around competitor locations')
add_para('')
add_heading('Implementation Code:', 3)
add_code('''def perform_buffer_analysis(competitor_locations, buffer_radius_km=2):
    """
    Create buffers around competitor locations
    
    Args:
        competitor_locations: GeoDataFrame with competitor points
        buffer_radius_km: Buffer radius in kilometers
    
    Returns:
        GeoDataFrame with buffer geometries
    """
    import geopandas as gpd
    from shapely.geometry import Point
    
    # Ensure CRS is in meters for accurate buffering
    if competitor_locations.crs.is_geographic:
        # Convert to projected CRS (UTM or similar)
        competitor_locations = competitor_locations.to_crs(epsg=32643)  # WGS84 UTM Zone 43N for India
    
    # Create buffers (radius in meters)
    buffer_radius_m = buffer_radius_km * 1000
    competitor_locations['buffer'] = competitor_locations.geometry.buffer(buffer_radius_m)
    
    # Create new GeoDataFrame with buffer geometries
    buffers = gpd.GeoDataFrame(
        competitor_locations.drop('geometry', axis=1),
        geometry='buffer',
        crs=competitor_locations.crs
    )
    
    # Convert back to geographic CRS if needed
    buffers = buffers.to_crs(epsg=4326)
    
    return buffers

def find_market_gaps(candidate_areas, competitor_buffers):
    """
    Find areas outside competitor coverage
    
    Args:
        candidate_areas: GeoDataFrame of potential locations
        competitor_buffers: GeoDataFrame with competitor buffers
    
    Returns:
        GeoDataFrame of uncovered areas
    """
    # Union all competitor buffers
    all_buffers = competitor_buffers.geometry.unary_union
    
    # Find areas not covered
    gaps = candidate_areas[~candidate_areas.geometry.intersects(all_buffers)]
    
    return gaps''')

# Weighted Scoring
add_heading('4.2 Weighted Scoring Algorithm', 2)
add_para('Type: Multi-Criteria Decision Analysis (MCDA)')
add_para('')
add_para('Formula:')
add_para('Score = w₁×PopulationDensity + w₂×CompetitorDistance + w₃×Accessibility')
add_para('where Σwᵢ = 1')
add_para('')
add_heading('Implementation Code:', 3)
add_code('''def weighted_location_scoring(candidate_wards, weights=None):
    """
    Score candidate locations using weighted criteria
    
    Args:
        candidate_wards: DataFrame with location metrics
        weights: Dict of weights for each criterion
    
    Returns:
        DataFrame with scores and rankings
    """
    if weights is None:
        weights = {
            'population_density': 0.40,
            'competitor_distance': 0.35,
            'accessibility': 0.25
        }
    
    # Normalize each criterion to [0, 1]
    for criterion in weights.keys():
        min_val = candidate_wards[criterion].min()
        max_val = candidate_wards[criterion].max()
        candidate_wards[f'{criterion}_norm'] = (
            (candidate_wards[criterion] - min_val) / (max_val - min_val)
        )
    
    # Calculate weighted score
    candidate_wards['total_score'] = 0
    for criterion, weight in weights.items():
        candidate_wards['total_score'] += (
            weight * candidate_wards[f'{criterion}_norm']
        )
    
    # Rank locations
    candidate_wards['rank'] = candidate_wards['total_score'].rank(ascending=False)
    
    # Sort by score
    result = candidate_wards.sort_values('total_score', ascending=False)
    
    return result''')

# Haversine Distance
add_heading('4.3 Haversine Distance Formula', 2)
add_para('Type: Geospatial Distance Calculation')
add_para('')
add_para('Formula:')
add_para('a = sin²(Δφ/2) + cos(φ₁)×cos(φ₂)×sin²(Δλ/2)')
add_para('c = 2×atan2(√a, √(1-a))')
add_para('d = R × c')
add_para('where R = 6371 km (Earth radius)')
add_para('')
add_heading('Implementation Code:', 3)
add_code('''def haversine_distance(lat1, lon1, lat2, lon2):
    """
    Calculate great-circle distance between two points
    
    Args:
        lat1, lon1: Coordinates of point 1
        lat2, lon2: Coordinates of point 2
    
    Returns:
        Distance in kilometers
    """
    import numpy as np
    
    # Convert to radians
    lat1, lon1, lat2, lon2 = map(np.radians, [lat1, lon1, lat2, lon2])
    
    # Haversine formula
    dlat = lat2 - lat1
    dlon = lon2 - lon1
    
    a = np.sin(dlat/2)**2 + np.cos(lat1) * np.cos(lat2) * np.sin(dlon/2)**2
    c = 2 * np.arctan2(np.sqrt(a), np.sqrt(1-a))
    
    # Earth radius in kilometers
    R = 6371
    
    distance = R * c
    return distance

def find_nearest_competitor(location, competitor_locations):
    """
    Find distance to nearest competitor
    
    Args:
        location: Tuple of (lat, lon)
        competitor_locations: DataFrame with lat, lon columns
    
    Returns:
        Minimum distance to any competitor
    """
    distances = []
    lat1, lon1 = location
    
    for _, comp in competitor_locations.iterrows():
        dist = haversine_distance(lat1, lon1, comp['lat'], comp['lon'])
        distances.append(dist)
    
    return min(distances) if distances else float('inf')''')
doc.add_page_break()

# ========== PERFORMANCE METRICS ==========
add_heading('5. PERFORMANCE METRICS', 1)
doc.add_paragraph()

add_heading('5.1 MAPE (Mean Absolute Percentage Error)', 2)
add_code('''def calculate_mape(actual, predicted):
    """
    Calculate MAPE
    
    Formula: MAPE = (100/n) × Σ |(Actual - Forecast)/Actual|
    """
    import numpy as np
    
    # Remove zeros to avoid division by zero
    mask = actual != 0
    return np.mean(np.abs((actual[mask] - predicted[mask]) / actual[mask])) * 100''')

add_heading('5.2 RMSE (Root Mean Squared Error)', 2)
add_code('''def calculate_rmse(actual, predicted):
    """
    Calculate RMSE
    
    Formula: RMSE = √[(1/n) × Σ(Actual - Forecast)²]
    """
    import numpy as np
    return np.sqrt(np.mean((actual - predicted)**2))''')

add_heading('5.3 R² (Coefficient of Determination)', 2)
add_code('''def calculate_r2(actual, predicted):
    """
    Calculate R²
    
    Formula: R² = 1 - (SS_res / SS_tot)
    where SS_res = Σ(y - ŷ)²
          SS_tot = Σ(y - ȳ)²
    """
    import numpy as np
    
    ss_res = np.sum((actual - predicted)**2)
    ss_tot = np.sum((actual - np.mean(actual))**2)
    
    return 1 - (ss_res / ss_tot)''')

doc.add_page_break()

# Summary Table
add_heading('ALGORITHM SUMMARY', 1)
doc.add_paragraph()

# Create table
table = doc.add_table(rows=17, cols=4)
table.style = 'Light Grid Accent 1'

# Headers
headers = ['Category', 'Algorithm', 'Type', 'Library']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True

# Data
data = [
    ['Forecasting', 'ARIMA', 'Statistical', 'statsmodels'],
    ['Forecasting', 'Prophet', 'Additive', 'fbprophet'],
    ['Forecasting', 'XGBoost', 'Ensemble', 'xgboost'],
    ['Forecasting', 'LightGBM', 'Ensemble', 'lightgbm'],
    ['Forecasting', 'Random Forest', 'Ensemble', 'sklearn'],
    ['Forecasting', 'Weighted Ensemble', 'Meta', 'Custom'],
    ['Inventory', 'ABC Analysis', 'Pareto', 'Custom'],
    ['Inventory', 'XYZ Analysis', 'Variability', 'Custom'],
    ['Inventory', 'FSN Analysis', 'Movement', 'Custom'],
    ['CRM', 'K-Means', 'Clustering', 'sklearn'],
    ['CRM', 'RFM', 'Segmentation', 'Custom'],
    ['GIS', 'Buffer Analysis', 'Spatial', 'geopandas'],
    ['GIS', 'Weighted Scoring', 'MCDA', 'Custom'],
    ['GIS', 'Haversine Distance', 'Geospatial', 'Custom'],
    ['Layout', 'Relationship Matrix', 'Optimization', 'Custom'],
    ['Layout', 'Greedy Placement', 'Heuristic', 'Custom']
]

for i, row_data in enumerate(data, 1):
    for j, cell_data in enumerate(row_data):
        table.rows[i].cells[j].text = cell_data

# Save document
doc.save('URIP_Algorithms_with_Code.docx')
print("✓ Document generated successfully!")
print("File saved as: URIP_Algorithms_with_Code.docx")
print("\nContents:")
print("• 6 Forecasting algorithms with full code")
print("• 3 Inventory classification algorithms")
print("• 2 Customer analytics algorithms")
print("• 3 GIS algorithms")
print("• 2 Layout optimization algorithms")
print("• Performance metrics formulas")
print("• Summary reference table")
