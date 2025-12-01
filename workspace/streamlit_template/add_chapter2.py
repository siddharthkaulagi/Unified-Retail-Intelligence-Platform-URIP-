"""Add Chapter 2 - Literature Survey with 10 detailed paragraphs"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

doc = Document('Final_Report_Phase_2.docx')

def format_para(p, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.alignment = align
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.bold = bold

def add_heading(text, level=1):
    if level == 1:
        p = doc.add_heading(text, level=1)
        format_para(p, 16, True, WD_ALIGN_PARAGRAPH.LEFT)
    else:
        p = doc.add_heading(text, level=2)
        format_para(p, 16, False, WD_ALIGN_PARAGRAPH.LEFT)
    return p

def add_para(text):
    p = doc.add_paragraph(text)
    format_para(p)
    return p

add_heading('CHAPTER 2: LITERATURE SURVEY', 1)
doc.add_paragraph()

lit_papers = [
    'Hu et al. (2025) [1] proposed an enhanced sales forecasting system integrating temporal pattern recognition with external market indicators for retail chains. Their approach combined convolutional neural networks for trend extraction with attention mechanisms to weight recent observations more heavily. Applied to multi-store grocery data, the model achieved fifteen percent improvement in mean absolute percentage error over baseline ARIMA implementations. The research demonstrated that incorporating promotional calendars and weather data as exogenous variables significantly enhanced prediction accuracy during high-volatility periods. This work directly informs our ensemble approach in URIP by validating the importance of multi-model strategies and external factor integration for robust retail forecasting.',
    
    'Venkatesh and Sowmiya (2024) [2] developed a hybrid inventory classification framework combining ABC analysis with machine learning clustering to identify optimal stock policies for different product segments. Their methodology used K-means clustering on velocity and value metrics followed by association rule mining to discover replenishment patterns. Implementation in a pharmaceutical retail chain reduced holding costs by eighteen percent while maintaining ninety-seven percent service levels. The study emphasized the inadequacy of single-dimensional classification in modern retail environments with diverse SKU characteristics. Our implementation of multi-criteria inventory classification in URIP draws heavily on their findings regarding the value of integrated classification approaches for strategic inventory management.',
    
    'Taparia (2024) [3] investigated the application of gradient boosting decision trees for short-term sales prediction in fast-moving consumer goods retail. The research compared XGBoost, LightGBM, and CatBoost algorithms across promotional and non-promotional periods using feature engineering techniques including lagged sales variables and categorical embeddings. LightGBM demonstrated superior performance with twenty-three percent lower RMSE compared to traditional time series methods while maintaining acceptable computational efficiency for real-time prediction scenarios. The analysis revealed that tree-based models excel in capturing non-linear relationships between promotional intensity and sales uplift. These insights guided our selection of XGBoost and LightGBM as core forecasting engines within URIP.',
    
    'Kandemir (2022) [4] presented a comprehensive framework for retail location analytics using geographic information systems integrated with demographic data and spatial optimization algorithms. The methodology employed Voronoi tessellation to define market boundaries, gravity modeling for catchment area estimation, and multi-criteria decision analysis incorporating distance decay functions. Application to a fashion retail expansion project identified twelve high-potential locations with projected revenue exceeding baseline predictions by thirty-one percent. The study highlighted the critical importance of competitor proximity analysis in site selection decisions. Our GIS module in URIP implements several concepts from this research including buffer-based competitive analysis and population-weighted scoring algorithms.',
    
    'Frank and Henry (2024) [5] explored ensemble learning strategies for demand forecasting in omnichannel retail environments. Their stacked generalization approach combined time series models, gradient boosting, and recurrent neural networks through a meta-learner trained on validation residuals. Testing across three retail categories demonstrated that ensemble methods consistently outperformed individual models with average accuracy improvements of fourteen to nineteen percent. The research emphasized the importance of model diversity in ensemble composition and appropriate weight assignment based on forecast horizon. These principles directly influenced our implementation of weighted ensemble forecasting in URIP where model contributions are optimized based on historical performance across different product categories.',
    
    'A comparative study of ARIMA, SARIMA, and LSTM for retail sales forecasting (2022) [6] benchmarked traditional statistical methods against deep learning approaches using twelve months of daily sales data from electronics retail. ARIMA proved effective for stationary series with clear linear trends. SARIMA captured seasonal patterns accurately for products with strong annual cycles. LSTM networks excelled in handling complex non-linear patterns but required substantial training data and computational resources. The analysis concluded that no single method universally dominates across all product categories and forecasting horizons. This finding validates our multi-model approach in URIP where different algorithms are evaluated and selected based on data characteristics and business requirements.',
    
    'Rana et al. (2023) [7] developed an integrated analytics platform for retail supply chain optimization incorporating forecasting, routing, and warehouse management modules. Their system architecture utilized microservices for modularity and scalability with a unified data layer supporting cross-module analytics. Implementation in a regional supermarket chain reduced stockout incidents by forty-two percent and improved delivery efficiency by twenty-six percent. The platform demonstrated that integrating multiple analytical capabilities within a single ecosystem yields synergistic benefits exceeding the sum of individual modules. This architecture philosophy strongly aligns with our design of URIP as a unified intelligence platform rather than a collection of standalone tools.',
    
    'Rani et al. (2024) [8] investigated customer segmentation methodologies for personalized marketing in retail using both traditional RFM analysis and advanced clustering algorithms. Their comparative study across three thousand customers revealed that DBSCAN clustering combined with RFM scores identified micro-segments with distinct behavioral patterns missed by conventional quintile-based segmentation. Targeted campaigns based on micro-segments achieved thirty-eight percent higher conversion rates compared to broad segment approaches. The research emphasized the value of granular segmentation for resource allocation in marketing budgets. Our CRM module integrates their recommended approach of combining rule-based RFM with data-driven clustering to support both strategic and tactical marketing decisions.',
    
    'Xu (2023) [9] presented a systematic methodology for retail facility layout optimization using relationship analysis and genetic algorithms. The approach quantified department affinities based on cross-purchase patterns and used evolutionary optimization to minimize total travel distance while respecting spatial constraints. Application to a department store redesign reduced average customer walking distance by twenty-nine percent and increased cross-category purchases by sixteen percent. The study validated that data-driven layout decisions significantly impact both customer experience and sales performance. This research directly informed our implementation of the facility planning module in URIP where relationship matrices are optimized through constraint satisfaction algorithms.',
    
    'Bauskar (2022) [10] examined the role of business intelligence dashboards in retail decision-making through user experience studies across multiple organizational roles. The research identified that dashboard effectiveness depends critically on appropriate visualization selection, hierarchical information architecture, and role-based customization. Managers prioritized high-level KPIs with drill-down capabilities while operational users required granular real-time metrics. Successful implementations balanced comprehensiveness with simplicity avoiding cognitive overload through progressive disclosure design patterns. These UX principles guided our dashboard development in URIP emphasizing intuitive navigation, contextual help, and export capabilities tailored to different user personas in retail organizations.'
]

for para in lit_papers:
    add_para(para)

doc.add_page_break()
doc.save('Final_Report_Phase_2.docx')
print("Chapter 2 (Literature Survey) added with 10 detailed paragraphs!")
