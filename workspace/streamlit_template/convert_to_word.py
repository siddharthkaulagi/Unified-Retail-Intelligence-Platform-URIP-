import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def parse_markdown_table(lines, doc):
    """Parses a markdown table and adds it to the document"""
    # Filter out separator lines (e.g. |---|---|)
    table_lines = [line for line in lines if not re.match(r'^\s*\|?[\s\-:|]+\|?\s*$', line)]
    
    if not table_lines:
        return

    # Determine dimensions
    rows = len(table_lines)
    cols = len(table_lines[0].split('|')) - 2 # Assuming | at start and end
    
    if cols < 1:
        cols = len(table_lines[0].split('|')) # Try without outer pipes

    try:
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        for i, line in enumerate(table_lines):
            # Split by pipe, remove first and last empty strings if they exist
            cells = [c.strip() for c in line.split('|')]
            if cells[0] == '': cells.pop(0)
            if cells[-1] == '': cells.pop(-1)
            
            row_cells = table.rows[i].cells
            for j, cell_text in enumerate(cells):
                if j < len(row_cells):
                    row_cells[j].text = cell_text.replace('**', '') # Remove bold markers
    except Exception as e:
        doc.add_paragraph(f"[Error rendering table: {e}]")
        for line in lines:
            doc.add_paragraph(line, style='No Spacing')

def convert_md_to_docx(md_file, docx_file):
    document = Document()
    
    # Set styles
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    for line in lines:
        line = line.strip()
        
        # Handle Code Blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                p = document.add_paragraph()
                p.style = 'No Spacing'
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_lines.append(line)
            continue

        # Handle Tables
        if line.startswith('|'):
            in_table = True
            table_lines.append(line)
            continue
        else:
            if in_table:
                parse_markdown_table(table_lines, document)
                table_lines = []
                in_table = False

        # Skip empty lines if not in block
        if not line:
            continue

        # Page Break
        if '<div style="page-break-after: always;"></div>' in line:
            document.add_page_break()
            continue

        # Headers
        if line.startswith('# '):
            document.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            document.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            document.add_heading(line[5:], level=4)
        
        # List Items
        elif line.startswith('* ') or line.startswith('- '):
            p = document.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line):
            p = document.add_paragraph(line, style='List Number')
            
        # Images (Placeholder)
        elif line.startswith('![') and '](' in line:
            alt_text = line[2:line.find(']')]
            document.add_paragraph(f"[IMAGE: {alt_text}]", style='Quote')
            
        # Normal Text
        else:
            p = document.add_paragraph()
            # Basic Bold formatting handling
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1: # Odd parts are between ** **
                    run.bold = True

    # Save
    document.save(docx_file)
    print(f"Successfully created {docx_file}")

if __name__ == "__main__":
    try:
        convert_md_to_docx('PROJECT_REPORT.md', 'PROJECT_REPORT.docx')
    except ImportError:
        print("python-docx not installed. Installing...")
        os.system('pip install python-docx')
        convert_md_to_docx('PROJECT_REPORT.md', 'PROJECT_REPORT.docx')
