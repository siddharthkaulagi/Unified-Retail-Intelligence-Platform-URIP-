"""
Generate College Project Report for URIP
Follows strict formatting guidelines
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_margins(doc, left=1.25, right=1, top=0.75, bottom=0.75):
    """Set page margins in inches"""
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)

def add_page_break(doc):
    """Add page break"""
    doc.add_page_break()

def set_paragraph_format(paragraph, font_size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, spacing=1.5):
    """Set paragraph formatting"""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = spacing
    paragraph_format.alignment = alignment
    
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(font_size)
        run.font.bold = bold

def add_chapter_heading(doc, text, level=1):
    """Add chapter heading with proper formatting"""
    if level == 1:
        p = doc.add_heading(text, level=1)
        set_paragraph_format(p, font_size=16, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.5)
    elif level == 2:
        p = doc.add_heading(text, level=2)
        set_paragraph_format(p, font_size=16, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.5)
    else:
        p = doc.add_heading(text, level=3)
        set_paragraph_format(p, font_size=14, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.5)
    return p

def add_body_paragraph(doc, text):
    """Add body paragraph with justified text"""
    p = doc.add_paragraph(text)
    set_paragraph_format(p, font_size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, spacing=1.5)
    return p

def generate_report():
    """Generate the complete project report"""
    doc = Document()
    
    # Set margins
    set_margins(doc)
    
    # ==================== OUTER TITLE PAGE ====================
    p = doc.add_paragraph('ML-DRIVEN UNIFIED RETAIL INTELLIGENCE PLATFORM (URIP)', style='Title')
    set_paragraph_format(p, font_size=18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
    
    doc.add_paragraph('\n' * 8)
    
    p = doc.add_paragraph('A Project Report')
    set_paragraph_format(p, font_size=14, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
    
    p = doc.add_paragraph('Submitted in partial fulfillment of the requirements for the award of the degree of')
    set_paragraph_format(p, font_size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
    
    p = doc.add_paragraph('BACHELOR OF ENGINEERING')
    set_paragraph_format(p, font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
    
    p = doc.add_paragraph('in')
    set_paragraph_format(p, font_size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
    
    p = doc.add_paragraph('COMPUTER SCIENCE AND ENGINEERING')
    set_paragraph_format(p, font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
    
    doc.add_paragraph('\n' * 6)
    
    p = doc.add_paragraph('[Your College Name]')
    set_paragraph_format(p, font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
    
    p = doc.add_paragraph('[Year]')
    set_paragraph_format(p, font_size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
    
    add_page_break(doc)
    
    # ==================== INNER TITLE PAGE ====================
    p = doc.add_paragraph('ML-DRIVEN UNIFIED RETAIL INTELLIGENCE PLATFORM (URIP)', style='Title')
    set_paragraph_format(p, font_size=18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
    
    doc.add_paragraph('\n' * 8)
    
    p = doc.add_paragraph('Submitted by')
    set_paragraph_format(p, font_size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
    
    p = doc.add_paragraph('[Student Name]')
    set_paragraph_format(p, font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
    
    p = doc.add_paragraph('[USN]')
    set_paragraph_format(p, font_size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
    
    doc.add_paragraph('\n' * 4)
    
    p = doc.add_paragraph('Under the guidance of')
    set_paragraph_format(p, font_size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
    
    p = doc.add_paragraph('[Guide Name]')
    set_paragraph_format(p, font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
    
    p = doc.add_paragraph('[Designation]')
    set_paragraph_format(p, font_size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
    
    doc.add_paragraph('\n' * 6)
    
    p = doc.add_paragraph('[Department of Computer Science and Engineering]')
    set_paragraph_format(p, font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
    
    p = doc.add_paragraph('[Your College Name]')
    set_paragraph_format(p, font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
    
    p = doc.add_paragraph('[Year]')
    set_paragraph_format(p, font_size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
    
    add_page_break(doc)
    
    # ==================== CERTIFICATE ====================
    p = doc.add_paragraph('CERTIFICATE', style='Title')
    set_paragraph_format(p, font_size=18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
    
    doc.add_paragraph('\n' * 2)
    
    cert_text = '''This is to certify that the project work entitled "ML-DRIVEN UNIFIED RETAIL INTELLIGENCE PLATFORM (URIP)" is a bonafide work carried out by [Student Name] (USN: [USN]) in partial fulfillment for the award of Bachelor of Engineering in Computer Science and Engineering of [University Name] during the academic year [Year].'''
    add_body_paragraph(doc, cert_text)
    
    doc.add_paragraph('\n' * 4)
    
    p = doc.add_paragraph('Internal Guide')
    set_paragraph_format(p, font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.5)
    
    p = doc.add_paragraph('[Guide Name]')
    set_paragraph_format(p, font_size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.5)
    
    p = doc.add_paragraph('[Designation]')
    set_paragraph_format(p, font_size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.5)
    
    doc.add_paragraph('\n' * 2)
    
    p = doc.add_paragraph('Head of Department')
    set_paragraph_format(p, font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.5)
    
    p = doc.add_paragraph('[HOD Name]')
    set_paragraph_format(p, font_size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.5)
    
    p = doc.add_paragraph('[Department]')
    set_paragraph_format(p, font_size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.5)
    
    add_page_break(doc)
    
    # ==================== ACKNOWLEDGMENT ====================
    p = doc.add_paragraph('ACKNOWLEDGMENT', style='Title')
    set_paragraph_format(p, font_size=18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
    
    doc.add_paragraph('\n')
    
    ack_text = '''I would like to express my sincere gratitude to all those who have contributed to the successful completion of this project. First and foremost, I am deeply grateful to my project guide, [Guide Name], for their invaluable guidance, continuous support, and insightful feedback throughout the development of this work. I would also like to thank the Head of the Department, [HOD Name], for providing the necessary facilities and support. My sincere thanks to the faculty members of the Department of Computer Science and Engineering for their suggestions and encouragement. I am also grateful to my family and friends for their constant support and motivation.'''
    add_body_paragraph(doc, ack_text)
    
    add_page_break(doc)
    
    # ==================== ABSTRACT ====================
    p = doc.add_paragraph('ABSTRACT', style='Title')
    set_paragraph_format(p, font_size=18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
    
    doc.add_paragraph('\n')
    
    abstract_text = '''The Unified Retail Intelligence Platform (URIP) is an integrated machine learning-driven system designed to address critical retail business challenges including sales forecasting, inventory optimization, and strategic location planning. The system employs multiple forecasting models such as ARIMA, Prophet, XGBoost, LightGBM, and Random Forest to predict future sales with high accuracy. It incorporates advanced inventory classification techniques (ABC, XYZ, FSN), GIS-based store location analysis, facility layout planning, and CRM analytics. The platform features an interactive dashboard built with Streamlit, providing comprehensive business intelligence through visualizations, automated report generation, and an AI-powered chatbot for natural language queries.'''
    add_body_paragraph(doc, abstract_text)
    
    add_page_break(doc)
    
    # ==================== TABLE OF CONTENTS ====================
    p = doc.add_paragraph('TABLE OF CONTENTS', style='Title')
    set_paragraph_format(p, font_size=18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
    
    doc.add_paragraph('\n')
    
    # Add TOC entries manually (auto-generation requires Word application)
    toc_entries = [
        ('CHAPTER 1: INTRODUCTION', '1'),
        ('1.1 Background', '1'),
        ('1.2 Motivation', '2'),
        ('1.3 Problem Description', '3'),
        ('1.4 Objectives', '4'),
        ('1.5 Scope', '5'),
        ('1.6 Applications', '6'),
        ('', ''),
        ('CHAPTER 2: LITERATURE SURVEY', '7'),
        ('', ''),
        ('CHAPTER 3: METHODOLOGY', '17'),
        ('', ''),
        ('CHAPTER 4: SYSTEM DESIGN', '25'),
        ('4.1 System Architecture', '25'),
        ('4.2 Module Descriptions', '27'),
        ('4.3 Use Case Diagram', '35'),
        ('4.4 Data Flow Diagram', '36'),
        ('', ''),
        ('CHAPTER 5: MODEL OVERVIEW', '37'),
        ('', ''),
        ('CHAPTER 6: IMPLEMENTATION', '50'),
        ('', ''),
        ('CHAPTER 7: RESULTS & DISCUSSION', '55'),
        ('', ''),
        ('CHAPTER 8: CONCLUSION & FUTURE WORK', '60'),
        ('', ''),
        ('REFERENCES', '62'),
    ]
    
    for entry, page in toc_entries:
        if entry:
            p = doc.add_paragraph(f'{entry}{"." * (80 - len(entry) - len(page))}{page}')
            set_paragraph_format(p, font_size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.5)
        else:
            doc.add_paragraph('')
    
    add_page_break(doc)
    
    # ==================== LIST OF FIGURES ====================
    p = doc.add_paragraph('LIST OF FIGURES', style='Title')
    set_paragraph_format(p, font_size=18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
    
    doc.add_paragraph('\n')
    
    figures = [
        ('Fig 1.1: URIP System Overview', '6'),
        ('Fig 3.1: Data Preprocessing Workflow', '18'),
        ('Fig 3.2: Forecasting Pipeline', '20'),
        ('Fig 3.3: GIS Analysis Workflow', '22'),
        ('Fig 4.1: System Architecture Diagram', '26'),
        ('Fig 4.2: Use Case Diagram', '35'),
        ('Fig 4.3: Data Flow Diagram (Level 0)', '36'),
        ('Fig 6.1: Dashboard Screenshot', '51'),
        ('Fig 6.2: GIS Map Visualization', '53'),
        ('Fig 7.1: Sales Forecast Comparison', '56'),
        ('Fig 7.2: Inventory Classification Results', '58'),
    ]
    
    for fig, page in figures:
        p = doc.add_paragraph(f'{fig}{"." * (80 - len(fig) - len(page))}{page}')
        set_paragraph_format(p, font_size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.5)
    
    add_page_break(doc)
    
    # ==================== LIST OF TABLES ====================
    p = doc.add_paragraph('LIST OF TABLES', style='Title')
    set_paragraph_format(p, font_size=18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
    
    doc.add_paragraph('\n')
    
    tables = [
        ('Table 3.1: Dataset Characteristics', '18'),
        ('Table 4.1: Module Comparison', '34'),
        ('Table 5.1: Model Comparison', '49'),
        ('Table 7.1: Model Performance Metrics', '57'),
    ]
    
    for table, page in tables:
        p = doc.add_paragraph(f'{table}{"." * (80 - len(table) - len(page))}{page}')
        set_paragraph_format(p, font_size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.5)
    
    add_page_break(doc)
    
    # ==================== CHAPTER 1: INTRODUCTION ====================
    add_chapter_heading(doc, 'CHAPTER 1: INTRODUCTION', level=1)
    doc.add_paragraph('')
    
    add_chapter_heading(doc, '1.1 Background', level=2)
    
    bg_text = '''The retail industry is experiencing a paradigm shift driven by digital transformation and data-driven decision-making. Modern retailers face complex challenges in managing inventory, predicting customer demand, optimizing store locations, and maintaining efficient supply chains. Traditional approaches to retail management often rely on manual processes, historical intuition, and reactive strategies that fall short in today's fast-paced, competitive marketplace. The exponential growth of transactional data combined with advances in machine learning and artificial intelligence has created unprecedented opportunities for retailers to gain actionable insights from their data. However, many retail businesses, especially small to medium enterprises, lack integrated platforms that can harness the power of machine learning for comprehensive business intelligence. The gap between data availability and actionable intelligence has motivated the development of unified platforms that can process historical sales data, forecast future trends, optimize inventory levels, analyze geographic market opportunities, and provide strategic recommendations—all within a single ecosystem. The Unified Retail Intelligence Platform (URIP) addresses this critical need by integrating multiple machine learning models, geographic information systems, inventory analytics, and customer relationship management into a cohesive, user-friendly application that empowers retail managers and decision-makers with cutting-edge analytical capabilities.'''
    add_body_paragraph(doc, bg_text)
    
    add_page_break(doc)
    
    add_chapter_heading(doc, '1.2 Motivation', level=2)
    
    mot_text = '''The motivation for developing URIP stems from several critical observations in the retail sector. First, sales forecasting remains one of the most challenging yet essential tasks for retail success. Inaccurate forecasts can lead to either stockouts resulting in lost sales opportunities or overstocking that ties up capital and increases storage costs. Second, inventory management continues to plague retailers, with optimal stock levels varying significantly across product categories, seasons, and store locations. Third, strategic location planning for new stores or distribution centers requires sophisticated analysis of demographic data, competitor locations, and accessibility factors—a capability typically reserved for large enterprises with dedicated analytics teams. Fourth, the proliferation of customer data has created opportunities for personalized marketing and customer segmentation, yet many retailers struggle to translate this data into actionable CRM strategies. Fifth, traditional business intelligence tools often operate in silos, requiring managers to consult multiple systems for comprehensive insights. Finally, the barrier to entry for machine learning adoption remains high due to the technical expertise required for model development and deployment. These challenges collectively motivated the creation of a unified platform that democratizes access to advanced retail analytics, making sophisticated forecasting, optimization, and planning capabilities accessible to businesses of all sizes without requiring deep technical expertise in data science or machine learning.'''
    add_body_paragraph(doc, mot_text)
    
    add_page_break(doc)
    
    add_chapter_heading(doc, '1.3 Problem Description', level=2)
    
    prob_text = '''Retail businesses encounter numerous interconnected problems that impact profitability and operational efficiency. The primary challenge is demand uncertainty—retailers must predict future sales with sufficient accuracy to make informed decisions about procurement, staffing, and promotions. However, sales data exhibits complex patterns including seasonal trends, promotional effects, external economic factors, and sudden market shifts that make forecasting extremely difficult. Traditional statistical methods like simple moving averages prove inadequate for capturing these nonlinear relationships. A second major problem is inventory optimization, where retailers must balance the competing objectives of minimizing holding costs while maximizing product availability. This requires sophisticated classification of products based on their value contribution, demand variability, and obsolescence rate. Third, geographic expansion decisions are often made based on intuition rather than data-driven analysis, leading to suboptimal store placements that fail to capture market potential or unnecessarily compete with existing locations. Fourth, retail facility layouts frequently evolve organically without systematic planning, resulting in inefficient material flows, poor space utilization, and customer navigation difficulties. Fifth, customer data remains underutilized, with segmentation and targeting strategies based on simple demographic rules rather than behavioral analytics. Sixth, managers lack real-time visibility into key performance indicators and must wait for periodic reports that may be outdated by the time they are reviewed. Finally, there is no single integrated system that addresses all these problems holistically, forcing businesses to cobble together disparate tools and manually integrate insights across different analytical workflows.'''
    add_body_paragraph(doc, prob_text)
    
    add_page_break(doc)
    
    add_chapter_heading(doc, '1.4 Objectives', level=2)
    
    obj_text = '''The primary objectives of the Unified Retail Intelligence Platform are as follows:'''
    add_body_paragraph(doc, obj_text)
    
    objectives = [
        'To develop an integrated machine learning platform that combines sales forecasting, inventory analytics, geographic intelligence, and customer relationship management into a unified solution.',
        'To implement and compare multiple state-of-the-art forecasting models including ARIMA, Prophet, XGBoost, LightGBM, and Random Forest, enabling automatic model selection based on performance metrics.',
        'To create a comprehensive inventory classification system using ABC analysis, XYZ analysis, and FSN analysis to support optimal stock management strategies.',
        'To integrate Geographic Information Systems (GIS) capabilities for analyzing store locations, identifying expansion opportunities, and visualizing market coverage.',
        'To provide facility layout planning tools that optimize department arrangements based on relationship analysis and material flow patterns.',
        'To implement customer segmentation using RFM (Recency, Frequency, Monetary) analysis and clustering techniques for targeted marketing campaigns.',
        'To develop an intuitive dashboard interface that visualizes key performance indicators, trends, and forecasts without requiring technical expertise from end users.',
        'To implement automated report generation capabilities that consolidate insights from multiple analytics modules into professional business reports.',
        'To integrate an AI-powered chatbot that can answer natural language queries about business data and provide instant insights.',
        'To ensure scalability and extensibility of the platform to accommodate future analytical modules and integration with external data sources.'
    ]
    
    for i, obj in enumerate(objectives, 1):
        p = doc.add_paragraph(f'{i}. {obj}', style='List Number')
        set_paragraph_format(p, font_size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, spacing=1.5)
    
    add_page_break(doc)
    
    add_chapter_heading(doc, '1.5 Scope', level=2)
    
    scope_text = '''The scope of URIP encompasses both technical and business dimensions. From a technical perspective, the platform is built using Python as the primary programming language, leveraging libraries such as scikit-learn, Prophet, XGBoost, LightGBM, GeoPandas, Folium, and Plotly for analytical and visualization capabilities. The user interface is developed using Streamlit, which provides a rapid application development framework for data science applications. The platform supports data ingestion from CSV and Excel files, with validation and preprocessing pipelines that handle missing values, outliers, and format inconsistencies. Forecasting capabilities extend up to customizable time horizons with confidence intervals and multiple model comparison. The GIS module supports geographical data in KML and CSV formats, enabling population analysis, competitive landscape mapping, and location recommendations. Inventory analytics covers three classification methodologies with visualizations and strategic recommendations. The CRM module implements both unsupervised clustering and supervised RFM segmentation with customer profiling capabilities. From a business scope perspective, URIP targets retail businesses across various sectors including grocery stores, supermarkets, specialty retail, and multi-location retail chains. The platform is designed for use by store managers, procurement officers, operations managers, marketing teams, and executive decision-makers. It handles transactional datasets ranging from thousands to millions of records, with scalability for expanding business operations. The platform provides insights across operational, tactical, and strategic time horizons—from daily sales forecasts to multi-year expansion planning. While the current implementation focuses on retail applications, the underlying architecture is designed to be adaptable to other sectors such as manufacturing, logistics, and service industries with minor customization.'''
    add_body_paragraph(doc, scope_text)
    
    add_page_break(doc)
    
    add_chapter_heading(doc, '1.6 Applications', level=2)
    
    app_text = '''The Unified Retail Intelligence Platform has numerous practical applications across the retail value chain. In demand forecasting and planning, retailers can generate accurate sales predictions to optimize procurement schedules, reducing instances of stockouts and minimizing excess inventory costs. The platform enables data-driven decisions about promotional campaigns by forecasting the impact of price changes and seasonal events. For inventory management, the ABC-XYZ-FSN classification system helps businesses prioritize their inventory control efforts, focusing resources on high-value items while automating replenishment for low-value products. This leads to improved working capital efficiency and better turnover rates. In the domain of geographic expansion and location intelligence, businesses can identify underserved markets by analyzing population density, competitor presence, and accessibility factors. The GIS capabilities support site selection for new stores, distribution centers, or warehouses by evaluating multiple location candidates against strategic criteria. For facility and layout optimization, the platform assists in designing efficient store layouts that minimize travel distances for high-traffic product pairs and optimize space utilization based on category sales contribution. In customer relationship management, businesses can segment their customer base for targeted marketing campaigns, identify high-value customers for loyalty programs, and detect at-risk customers who may require retention interventions. The automated reporting feature enables routine generation of business intelligence reports for stakeholders, board meetings, and regulatory compliance. The AI chatbot provides instant access to business insights for managers who need quick answers without navigating through complex dashboards. Educational institutions can use URIP as a teaching tool for courses in business analytics, data science, and retail management. Small and medium enterprises particularly benefit from URIP as it provides enterprise-grade analytics capabilities without the need for expensive commercial software or dedicated data science teams.'''
    add_body_paragraph(doc, app_text)
    
    doc.add_paragraph('\n')
    add_body_paragraph(doc, '<Insert Fig 1.1: URIP System Overview Here>')
    
    add_page_break(doc)
    
    # Save the document
    output_path = 'URIP_Project_Report_Part1.docx'
    doc.save(output_path)
    return output_path

if __name__ == '__main__':
    path = generate_report()
    print(f'Report Part 1 generated: {path}')
