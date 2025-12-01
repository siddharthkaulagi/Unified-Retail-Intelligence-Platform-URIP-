import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import numpy as np
import networkx as nx
from datetime import datetime
import io
import base64
import requests
import json
import tempfile
import os

from utils.ui_components import render_sidebar


# --- PATHS ---
# Get the absolute path of the directory containing the script's parent
BASE_DIR = os.path.dirname(os.path.dirname(os.path.realpath(__file__)))
ASSETS_DIR = os.path.join(BASE_DIR, "assets")
CSS_FILE = os.path.join(ASSETS_DIR, "custom.css")


def load_css():
    if os.path.exists(CSS_FILE):
        with open(CSS_FILE) as f:
            st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)

load_css()

# Optional libraries for image export
try:
    from PIL import Image as PILImage
    pillow_available = True
except Exception:
    pillow_available = False

try:
    from docx import Document
    from docx.shared import Inches
    docx_available = True
except Exception:
    docx_available = False

# Import Gemini AI assistant
try:
    from utils.gemini_ai import get_gemini_assistant, analyze_forecast_file, GEMINI_AVAILABLE
    claude_available = GEMINI_AVAILABLE
except Exception as e:
    claude_available = False
    st.error(f"❌ Gemini AI not available: {str(e)}")
    st.info("💡 **To enable Gemini AI features:** `pip install google-generativeai`")

# Initialize AI analysis result variable
ai_analysis_result = None





def get_ai_recommendations(forecast_data, business_context):
    """Generate AI-driven recommendations using Gemini AI"""
    if not claude_available:
        return generate_fallback_recommendations(forecast_data)
    
    try:
        # Initialize Gemini AI assistant
        from utils.gemini_ai import GeminiChatbot
        chatbot = GeminiChatbot()
        
        # Create detailed prompt for personalized recommendations
        prompt = f"""
        You are a senior retail analytics consultant. Based on the following sales forecast data and business context, provide 7 SPECIFIC and ACTIONABLE business recommendations.
        
        FORECAST DATA:
        - Total Forecasted Sales: ₹{forecast_data.get('total_forecast', 0):,.0f}
        - Average Daily Sales: ₹{forecast_data.get('avg_daily', 0):,.0f}
        - Best Performing Model: {forecast_data.get('best_model', 'N/A')}
        - Forecast Period: {forecast_data.get('period', 30)} days
        - Demand Volatility: {forecast_data.get('volatility', 'Moderate')}
        - Growth Rate: {forecast_data.get('growth_rate', 0):.1f}%
        
        BUSINESS CONTEXT:
        {business_context}
        
        Provide exactly 7 recommendations in this format (one per line):
        1. [Category]: [Specific actionable recommendation based on the actual numbers provided]
        
        Categories should include: Inventory Management, Staff Planning, Financial Planning, Sales Strategy, Target Setting, Performance Monitoring, Supply Chain
        
        Make each recommendation SPECIFIC to the data provided above. Reference actual numbers, growth rates, and volatility levels.
        Keep each recommendation to 1-2 sentences maximum.
        """
        
        # Get AI response
        response = chatbot.chat(prompt)
        
        if response.get('success'):
            # Parse AI response into list of recommendations
            ai_text = response['response']
            recommendations = []
            
            # Split by numbered lines
            lines = ai_text.strip().split('\n')
            for line in lines:
                line = line.strip()
                # Look for lines that start with numbers
                if line and (line[0].isdigit() or line.startswith('-') or line.startswith('•')):
                    # Clean up the line
                    cleaned = line.lstrip('0123456789.-•) ').strip()
                    if cleaned:
                        # Add icon prefix based on category
                        if 'inventory' in cleaned.lower():
                            recommendations.append(f"📦 {cleaned}")
                        elif 'staff' in cleaned.lower() or 'workforce' in cleaned.lower():
                            recommendations.append(f"👥 {cleaned}")
                        elif 'financial' in cleaned.lower() or 'budget' in cleaned.lower():
                            recommendations.append(f"💰 {cleaned}")
                        elif 'sales' in cleaned.lower():
                            recommendations.append(f"📈 {cleaned}")
                        elif 'target' in cleaned.lower():
                            recommendations.append(f"🎯 {cleaned}")
                        elif 'performance' in cleaned.lower() or 'monitor' in cleaned.lower():
                            recommendations.append(f"📊 {cleaned}")
                        elif 'supply' in cleaned.lower():
                            recommendations.append(f"🔄 {cleaned}")
                        else:
                            recommendations.append(f"💡 {cleaned}")
            
            # Return AI recommendations if we got at least 5
            if len(recommendations) >= 5:
                return recommendations[:7]  # Return max 7
        
        # Fallback if AI fails
        return generate_fallback_recommendations(forecast_data)
        
    except Exception as e:
        logger.error(f"Error getting AI recommendations: {str(e)}")
        return generate_fallback_recommendations(forecast_data)

def generate_fallback_recommendations(forecast_data):
    """Generate fallback recommendations when AI is unavailable"""
    recommendations = [
        f"📦 **Inventory Management**: Based on ₹{forecast_data.get('total_forecast', 0):,.0f} forecasted sales, maintain optimal stock levels to avoid overstocking",
        f"👥 **Staff Planning**: Schedule {forecast_data.get('staff_suggestion', 'adequate')} staff during peak sales periods",
        f"💰 **Financial Planning**: Prepare budget for ₹{forecast_data.get('total_forecast', 0):,.0f} projected revenue with 10-15% buffer for uncertainty",
        f"📈 **Sales Strategy**: Focus on high-performing categories that show consistent demand patterns",
        f"🎯 **Target Setting**: Set realistic sales targets based on {forecast_data.get('volatility', 'moderate')} demand volatility",
        f"📊 **Performance Monitoring**: Track actual vs forecasted sales weekly and adjust operations accordingly",
        f"🔄 **Supply Chain**: Review supplier relationships to ensure reliable delivery for forecasted demand"
    ]
    return recommendations

def analyze_business_metrics(data_df, forecast_results):
    """Analyze business metrics and generate insights"""
    insights = {}

    try:
        # Sales analysis
        if 'Sales' in data_df.columns:
            total_sales = data_df['Sales'].sum()
            avg_daily_sales = data_df['Sales'].mean()
            sales_volatility = data_df['Sales'].std() / avg_daily_sales if avg_daily_sales > 0 else 0

            insights['total_sales'] = total_sales
            insights['avg_daily'] = avg_daily_sales
            insights['volatility'] = sales_volatility

            # Category analysis
            if 'Category' in data_df.columns:
                category_performance = data_df.groupby('Category')['Sales'].agg(['sum', 'mean', 'count']).round(2)
                top_category = category_performance.loc[category_performance['sum'].idxmax()]
                insights['top_category'] = top_category.to_dict()

            # Store analysis
            if 'Store' in data_df.columns:
                store_performance = data_df.groupby('Store')['Sales'].agg(['sum', 'mean', 'count']).round(2)
                best_store = store_performance.loc[store_performance['sum'].idxmax()]
                insights['best_store'] = best_store.to_dict()

        # Forecast analysis
        best_model = None
        best_mae = float('inf')

        for model_name, result in forecast_results.items():
            if 'metrics' in result:
                mae = result['metrics'].get('mae', float('inf'))
                if mae < best_mae:
                    best_mae = mae
                    best_model = model_name

        insights['best_model'] = best_model
        insights['best_mae'] = best_mae

        # Generate forecast metrics
        if best_model and 'forecast' in forecast_results[best_model]:
            forecast_df = forecast_results[best_model]['forecast']
            total_forecast = forecast_df['yhat'].sum()
            avg_forecast = forecast_df['yhat'].mean()

            insights['total_forecast'] = total_forecast
            insights['avg_forecast'] = avg_forecast

            # Calculate growth rate
            if insights['total_sales'] > 0:
                growth_rate = ((total_forecast - insights['total_sales']) / insights['total_sales']) * 100
                insights['growth_rate'] = growth_rate

    except Exception as e:
        st.warning(f"Error in business analysis: {str(e)}")

    return insights

st.set_page_config(page_title="Reports", page_icon="📋", layout="wide")

# Check authentication
if not st.session_state.get('authenticated', False):
    st.error("Please login first!")
    st.stop()

# Sidebar
# Sidebar
render_sidebar()

st.title("📋 Reports & Downloads")
st.markdown("Generate comprehensive reports and download forecasting results")

# Check if forecast results are available
if 'forecast_results' not in st.session_state:
    st.warning("⚠️ No forecast results found. Please run forecasting first!")
    st.markdown("[👆 Go to Model Selection page](2_🔮_Model_Selection)")
    st.stop()

results = st.session_state.forecast_results
config = st.session_state.get('forecast_config', {})



# Report Configuration
st.markdown("### ⚙️ Report Configuration")

report_sections = st.multiselect(
    "Report Sections",
    ["Executive Summary", "Technical Details", "Business Insights", "Visualizations", "Category Analysis", "Recommendations"],
    default=[],
    help="Select sections to include in the report"
)


# Additional Report Modules
st.markdown("### 🧩 Additional Modules")
mod_col1, mod_col2 = st.columns(2)

with mod_col1:
    include_facility = st.checkbox("Include Facility Layout", value=False)
    facility_options = []
    if include_facility:
        facility_options = st.multiselect(
            "Facility Components",
            ["Optimized Layout", "Relationship Network", "Material Flow"],
            default=["Optimized Layout", "Relationship Network", "Material Flow"]
        )

with mod_col2:
    include_gis = st.checkbox("Include Store Location GIS", value=False)
    gis_options = []
    if include_gis:
        gis_options = st.multiselect(
            "GIS Components",
            ["Network Statistics", "Store Locations", "Strategic Recommendations"],
            default=["Network Statistics", "Store Locations", "Strategic Recommendations"]
        )

# Let user choose which columns represent date/value/category for reports
# Build a list of candidate columns from uploaded data or model outputs
col_candidates = set()
uploaded = st.session_state.get('uploaded_data')
if uploaded is not None and isinstance(uploaded, pd.DataFrame):
    col_candidates.update(list(uploaded.columns))
else:
    # look into first available model data for column names
    for r in results.values():
        if isinstance(r, dict):
            if 'historical' in r and isinstance(r['historical'], pd.DataFrame):
                col_candidates.update(list(r['historical'].columns))
            if 'forecast' in r and isinstance(r['forecast'], pd.DataFrame):
                col_candidates.update(list(r['forecast'].columns))

col_candidates = list(col_candidates)
if not col_candidates:
    col_candidates = ['ds', 'y', 'yhat', 'Category', 'Sales']

date_col_default = config.get('date_col', 'ds') if config.get('date_col') in col_candidates else ('ds' if 'ds' in col_candidates else col_candidates[0])
value_col_default = config.get('value_col', 'y') if config.get('value_col') in col_candidates else ('y' if 'y' in col_candidates else col_candidates[0])
category_col_default = config.get('category_col', 'Category') if config.get('category_col') in col_candidates else ('Category' if 'Category' in col_candidates else (col_candidates[0] if col_candidates else 'Category'))

colA, colB, colC = st.columns(3)
with colA:
    date_column = st.selectbox('Date column for reports', col_candidates, index=col_candidates.index(date_col_default) if date_col_default in col_candidates else 0)
with colB:
    value_column = st.selectbox('Value column for reports (historical)', col_candidates, index=col_candidates.index(value_col_default) if value_col_default in col_candidates else 0)
with colC:
    category_column = st.selectbox('Category column for reports', col_candidates, index=col_candidates.index(category_col_default) if category_col_default in col_candidates else 0)

# persist into config so generate_docx_bytes can read them
config['date_col'] = date_column
config['value_col'] = value_column
config['category_col'] = category_column

# Report Preview
st.markdown("### 📄 Report Preview")
st.caption("Preview sections will be included in the downloaded Word report. Click to expand.")

# Executive Summary Expander
if "Executive Summary" in report_sections:
    with st.expander("📊 Executive Summary", expanded=True):
        # Key findings
        best_model = None
        best_mae = float('inf')
        total_forecast = 0
        
        for model_name, result in results.items():
            if 'metrics' in result:
                mae = result['metrics'].get('mae', float('inf'))
                if mae < best_mae:
                    best_mae = mae
                    best_model = model_name
            
            if 'forecast' in result:
                total_forecast += result['forecast']['yhat'].sum()
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Best Model", best_model or "N/A")
        with col2:
            st.metric("Best MAE", f"{best_mae:.2f}" if best_mae != float('inf') else "N/A")
        with col3:
            st.metric("Forecast Period", f"{config.get('forecast_days', 'N/A')} days")
        with col4:
            st.metric("Total Forecast", f"₹{total_forecast:,.2f}")

# Technical Details Expander
if "Technical Details" in report_sections:
    with st.expander("🔬 Technical Analysis"):
        metrics_data = []
        for model_name, result in results.items():
            if 'metrics' in result:
                metrics = result['metrics']
                metrics_data.append({
                    'Model': model_name,
                    'MAE': round(metrics.get('mae', 0), 4),
                    'RMSE': round(metrics.get('rmse', 0), 4),
                    'MAPE': round(metrics.get('mape', 0), 4) if metrics.get('mape') is not None else None,
                    'R2': round(metrics.get('r2', float('nan')), 4)
                })
        
        if metrics_data:
            metrics_df = pd.DataFrame(metrics_data)
            st.dataframe(metrics_df, use_container_width=True)

# Business Insights Expander  
if "Business Insights" in report_sections:
    with st.expander("📊 Business Intelligence"):
        # Get data for analysis
        data_df = st.session_state.get('uploaded_data')
        if data_df is None:
            data_df = st.session_state.get('processed_data')
        
        if data_df is not None:
            # Analyze business metrics
            insights = analyze_business_metrics(data_df, results)

            # Business impact analysis
            st.markdown("**Business Impact Analysis**")

            # Compute best_model for this section
            best_model_bi = None
            best_mae_bi = float('inf')
            for model_name, result in results.items():
                if 'metrics' in result:
                    mae = result['metrics'].get('mae', float('inf'))
                    if mae < best_mae_bi:
                        best_mae_bi = mae
                        best_model_bi = model_name

            if best_model_bi and 'forecast' in results[best_model_bi]:
                forecast_df = results[best_model_bi]['forecast']

                # Calculate business metrics
                total_forecast = forecast_df['yhat'].sum()
                avg_daily_sales = forecast_df['yhat'].mean()

                col1, col2 = st.columns(2)

                with col1:
                    st.markdown("**Revenue Projections**")
                    st.write(f"• Total forecasted: ₹{total_forecast:,.2f}")
                    st.write(f"• Daily average: ₹{avg_daily_sales:,.2f}")

                with col2:
                    st.markdown("**Operational Insights**")
                    cv = (forecast_df['yhat'].std() / forecast_df['yhat'].mean()) * 100
                    volatility = "Low" if cv < 10 else "Moderate" if cv < 20 else "High"
                    st.write(f"• Sales volatility: {volatility}")
                    
                    # Weekly pattern
                    forecast_df_copy = forecast_df.copy()
                    forecast_df_copy['day_of_week'] = pd.to_datetime(forecast_df_copy['ds']).dt.day_name()
                    weekly_avg = forecast_df_copy.groupby('day_of_week')['yhat'].mean()
                    best_day = weekly_avg.idxmax()
                    st.write(f"• Best sales day: {best_day}")

# Visualizations Section (separate from preview)
if "Visualizations" in report_sections:
    st.markdown("### 📈 Business Visualizations")
    st.caption("These charts will be embedded in the Word report")
    
    data_df = st.session_state.get('uploaded_data')
    if data_df is None:
        data_df = st.session_state.get('processed_data')
    
    if data_df is not None:
        col1, col2, col3 = st.columns(3)

        with col1:
            # Sales box plot
            if 'Sales' in data_df.columns:
                fig_box = px.box(data_df, y='Sales', title="Sales Distribution Box Plot")
                fig_box.update_layout(template='plotly_white')
                st.plotly_chart(fig_box, use_container_width=True)

        with col2:
            # Day of week pattern
            if 'Date' in data_df.columns and 'Sales' in data_df.columns:
                daily_sales_temp = data_df.copy()
                daily_sales_temp['Date'] = pd.to_datetime(daily_sales_temp['Date'], dayfirst=True, errors='coerce')
                daily_sales_temp['Day_of_Week'] = daily_sales_temp['Date'].dt.day_name()
                weekly_pattern = daily_sales_temp.groupby('Day_of_Week')['Sales'].mean().reindex(
                    ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
                )

                fig_weekly = px.bar(x=weekly_pattern.index, y=weekly_pattern.values,
                                  title="Average Sales by Day of Week",
                                  labels={'x': 'Day', 'y': 'Average Sales (₹)'})
                fig_weekly.update_layout(template='plotly_white')
                st.plotly_chart(fig_weekly, use_container_width=True)

        with col3:
            # Monthly seasonality
            if 'Date' in data_df.columns and 'Sales' in data_df.columns:
                daily_sales_temp = data_df.copy()
                daily_sales_temp['Date'] = pd.to_datetime(daily_sales_temp['Date'], dayfirst=True, errors='coerce')
                
                if len(daily_sales_temp) >= 60:  # At least 2 months
                    daily_sales_temp['Month'] = daily_sales_temp['Date'].dt.month_name()
                    monthly_pattern = daily_sales_temp.groupby('Month')['Sales'].mean()

                    fig_monthly = px.bar(x=monthly_pattern.index, y=monthly_pattern.values,
                                       title="Average Sales by Month",
                                       labels={'x': 'Month', 'y': 'Average Sales (₹)'})
                    fig_monthly.update_layout(template='plotly_white')
                    st.plotly_chart(fig_monthly, use_container_width=True)
                else:
                    st.info("Need at least 2 months of data for monthly pattern analysis")


        # AI-Driven Recommendations Section
        st.markdown("#####Business Recommendations")

        # Prepare forecast data for AI
        forecast_data = {
            'total_forecast': insights.get('total_forecast', 0),
            'avg_daily': insights.get('avg_forecast', 0),
            'best_model': insights.get('best_model', 'N/A'),
            'period': config.get('forecast_days', 30),
            'volatility': 'High' if insights.get('volatility', 0) > 0.3 else 'Moderate' if insights.get('volatility', 0) > 0.15 else 'Low',
            'growth_rate': insights.get('growth_rate', 0)
        }

        # Business context
        business_context = f"""
        Business Type: Retail Store
        Data Period: {len(data_df)} transactions
        Categories: {data_df['Category'].nunique() if 'Category' in data_df.columns else 'N/A'} product categories
        Stores: {data_df['Store'].nunique() if 'Store' in data_df.columns else 'Single store'}
        Average Transaction Value: ₹{insights.get('avg_daily', 0):,.0f}
        Sales Volatility: {forecast_data['volatility']}
        Forecasted Growth: {forecast_data['growth_rate']:.1f}%
        """

        # Get AI recommendations
        with st.spinner("🤖 Generating AI-powered recommendations..."):
            ai_recommendations = get_ai_recommendations(forecast_data, business_context)

        # Display recommendations
        st.markdown("**💡 Strategic Recommendations:**")
        for i, rec in enumerate(ai_recommendations[:7], 1):  # Show top 7 recommendations
            # Clean up recommendation text for display (remove excessive icons)
            clean_rec = rec.replace('📦', '').replace('👥', '').replace('💰', '').replace('📈', '').replace('🎯', '').replace('📊', '').replace('🔄', '')
            clean_rec = clean_rec.replace('**', '')  # Remove bold formatting
            st.write(f"{i}. {clean_rec.strip()}")

        # Pareto Chart - 80/20 Analysis
        st.markdown("##### 📊 Pareto Analysis (80/20 Rule)")
        st.caption("Identify which products/categories contribute to 80% of sales")
        
        if 'Category' in data_df.columns and 'Sales' in data_df.columns:
            category_sales = data_df.groupby('Category')['Sales'].sum().sort_values(ascending=False).reset_index()
            category_sales['Cumulative %'] = (category_sales['Sales'].cumsum() / category_sales['Sales'].sum()) * 100
            
            fig_pareto = go.Figure()
            
            # Bars
            fig_pareto.add_trace(go.Bar(
                x=category_sales['Category'],
                y=category_sales['Sales'],
                name='Sales',
                marker_color='lightblue',
                yaxis='y'
            ))
            
            # Cumulative line
            fig_pareto.add_trace(go.Scatter(
                x=category_sales['Category'],
                y=category_sales['Cumulative %'],
                name='Cumulative %',
                marker_color='red',
                yaxis='y2',
                mode='lines+markers',
                line=dict(width=2)
            ))
            
            # Add 80% reference line
            fig_pareto.add_hline(y=80, line_dash="dash", line_color="green", yref='y2',
                                annotation_text="80% Target", annotation_position="right")
            
            fig_pareto.update_layout(
                title="Pareto Chart: Category Sales Contribution",
                xaxis_title="Category",
                yaxis=dict(title="Sales (₹)"),
                yaxis2=dict(title="Cumulative %", overlaying='y', side='right', range=[0, 100]),
                template='plotly_white',
                height=400
            )
            
            st.plotly_chart(fig_pareto, use_container_width=True)
            
            # Insights
            top_20_pct_categories = int(len(category_sales) * 0.2) or 1
            top_contributors = category_sales.head(top_20_pct_categories)
            contribution = top_contributors['Sales'].sum() / category_sales['Sales'].sum() * 100
            
            st.info(f"💡 **Insight**: Top {top_20_pct_categories} categories ({top_20_pct_categories/len(category_sales)*100:.0f}%) contribute {contribution:.1f}% of total sales")

        # Correlation Matrix
        viz_tab1, viz_tab2 = st.tabs(["📈 Trends", "🔗 Correlations"])
        
        with viz_tab2:
            st.markdown("##### Correlation Matrix")
            st.caption("Understand relationships between numerical variables")
            
            numeric_cols = data_df.select_dtypes(include=[np.number]).columns.tolist()
            if len(numeric_cols) >= 2:
                corr_matrix = data_df[numeric_cols].corr()
                
                fig_corr = px.imshow(
                    corr_matrix,
                    text_auto='.2f',
                    aspect='auto',
                    color_continuous_scale='RdBu_r',
                    color_continuous_midpoint=0
                )
                
                fig_corr.update_layout(
                    title="Variable Correlation Matrix",
                    height=500
                )
                
                st.plotly_chart(fig_corr, use_container_width=True)
                
                st.caption("💡 Red = Positive correlation, Blue = Negative correlation, White = No correlation")
        
        with viz_tab1:
            st.markdown("##### 🔥 Sales Heatmap: Temporal Patterns")
            st.caption("Identify peak sales periods by day and month")
            
            if 'Date' in data_df.columns and 'Sales' in data_df.columns:
                heat_df = data_df.copy()
                heat_df['Date'] = pd.to_datetime(heat_df['Date'], dayfirst=True, errors='coerce')
                heat_df['Day_of_Week'] = heat_df['Date'].dt.day_name()
                heat_df['Month'] = heat_df['Date'].dt.month_name()
                
                # Create pivot table
                heatmap_data = heat_df.groupby(['Day_of_Week', 'Month'])['Sales'].sum().reset_index()
                heatmap_pivot = heatmap_data.pivot(index='Day_of_Week', columns='Month', values='Sales')
                
                # Reorder rows and columns
                day_order = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
                month_order = ['January', 'February', 'March', 'April', 'May', 'June', 
                              'July', 'August', 'September', 'October', 'November', 'December']
                
                heatmap_pivot = heatmap_pivot.reindex(day_order, fill_value=0)
                heatmap_pivot = heatmap_pivot[[col for col in month_order if col in heatmap_pivot.columns]]
                
                fig_heat = px.imshow(
                    heatmap_pivot,
                    labels=dict(x="Month", y="Day of Week", color="Sales (₹)"),
                    x=heatmap_pivot.columns,
                    y=heatmap_pivot.index,
                    color_continuous_scale='Reds',
                    aspect='auto'
                )
                
                fig_heat.update_layout(
                    title="Sales Heatmap: Day of Week vs Month",
                    height=400
                )
                
                st.plotly_chart(fig_heat, use_container_width=True)
                
                # Insights
                max_sales_day = heatmap_pivot.sum(axis=1).idxmax()
                max_sales_month = heatmap_pivot.sum(axis=0).idxmax()
                st.info(f"💡 **Peak Periods**: {max_sales_day} is the busiest day, {max_sales_month} is the busiest month")


        # Actionable insights based on data
        st.markdown("##### 🎯 Actionable Insights")

        col1, col2, col3 = st.columns(3)

        with col1:
            st.markdown("**📦 Inventory Optimization**")
            if 'top_category' in insights:
                top_cat = insights['top_category']
                st.info(f"Focus on {top_cat.get('Category', 'top')} category - represents {top_cat.get('sum', 0):.1f}% of sales")

            if forecast_data['volatility'] == 'High':
                st.warning("High volatility detected - maintain higher safety stock levels")
            else:
                st.success("Stable demand pattern - optimize inventory turnover")

        with col2:
            st.markdown("**👥 Staffing Strategy**")
            if 'best_store' in insights:
                best_st = insights['best_store']
                st.info(f"Allocate more staff to {best_st.get('Store', 'top')} store during peak hours")

            # Staff suggestion based on forecast
            if forecast_data['avg_daily'] > insights.get('avg_daily', 0) * 1.2:
                staff_suggestion = "additional"
            else:
                staff_suggestion = "standard"
            st.info(f"Staff planning: Schedule {staff_suggestion} staff for forecasted period")

        with col3:
            st.markdown("**💰 Financial Planning**")
            if 'growth_rate' in insights:
                growth = insights['growth_rate']
                if growth > 10:
                    st.success(f"Positive growth trend ({growth:.1f}%) - invest in expansion")
                elif growth < -5:
                    st.warning(f"Declining trend ({growth:.1f}%) - focus on cost optimization")
                else:
                    st.info(f"Stable growth ({growth:.1f}%) - maintain current operations")

            # Budget recommendation
            budget_buffer = 15 if forecast_data['volatility'] == 'High' else 10
            st.info(f"Budget planning: Allocate {budget_buffer}% buffer for demand uncertainty")

        # Performance alerts
        st.markdown("##### 🚨 Performance Alerts")

        alert_col1, alert_col2 = st.columns(2)

        with alert_col1:
            # Volatility alert
            if insights.get('volatility', 0) > 0.3:
                st.error("⚠️ **High Sales Volatility**: Consider diversifying product portfolio")

            # Low performing categories
            if 'Category' in data_df.columns and 'Sales' in data_df.columns:
                category_sales_summary = data_df.groupby('Category')['Sales'].sum()
                low_threshold = category_sales_summary.quantile(0.25)
                low_performers = category_sales_summary[category_sales_summary < low_threshold]
                if len(low_performers) > 0:
                    st.warning(f"📉 **Low Performing Categories**: {len(low_performers)} categories need attention")

        with alert_col2:
            # Growth alert
            if 'growth_rate' in insights and insights['growth_rate'] < -10:
                st.error("📉 **Negative Growth Trend**: Implement immediate corrective actions")

            # Model accuracy alert
            if insights.get('best_mae', float('inf')) > insights.get('avg_daily', 0) * 0.2:
                st.warning("🎯 **Model Accuracy**: Consider retraining models with more data")


# Generate downloadable reports
st.markdown("### 📥 Download Reports")

col1, col2, col3 = st.columns(3)

# Forecast vs Actual Comparison Table (Academic Format)
st.markdown("### 📊 Forecast vs Actual Comparison")
st.markdown("Detailed comparison of historical and forecasted values for model validation")

# Get best model by MAE
best_model_name = None
best_mae_value = float('inf')
for model_name, model_result in results.items():
    if 'metrics' in model_result:
        mae = model_result['metrics'].get('mae', None)
        try:
            if mae is not None and float(mae) < best_mae_value:
                best_mae_value = float(mae)
                best_model_name = model_name
        except Exception:
            pass

if best_model_name and best_model_name in results:
    best_result = results[best_model_name]
    
    st.info(f"**Best Performing Model:** {best_model_name} (MAE: {best_mae_value:.2f})")
    
    # Create comparison table
    comparison_data = []
    
    # Get historical data (Actual values)
    if 'historical' in best_result:
        historical_df = best_result['historical'].copy()
        
        # Format dates - remove timestamps
        if 'ds' in historical_df.columns:
            historical_df['ds'] = pd.to_datetime(historical_df['ds'])
            historical_df['Date'] = historical_df['ds'].dt.strftime('%Y-%m-%d')
        
        # Get last 30 days of historical data
        historical_df = historical_df.tail(30)
        
        for idx, row in historical_df.iterrows():
            comparison_data.append({
                'Sr.': len(comparison_data) + 1,
                'Date': row.get('Date', row.get('ds', '')),
                'Actual': f"{row.get('y', 0):,.2f}",
                'Type': 'Historical'
            })
    
    # Get forecast data (Predicted values)
    if 'forecast' in best_result:
        forecast_df = best_result['forecast'].copy()
        
        # Format dates - remove timestamps
        if 'ds' in forecast_df.columns:
            forecast_df['ds'] = pd.to_datetime(forecast_df['ds'])
            forecast_df['Date'] = forecast_df['ds'].dt.strftime('%Y-%m-%d')
        
        # Get first 30 days of forecast
        forecast_df = forecast_df.head(30)
        
        for idx, row in forecast_df.iterrows():
            comparison_data.append({
                'Sr.': len(comparison_data) + 1,
                'Date': row.get('Date', row.get('ds', '')),
                'Forecast': f"{row.get('yhat', 0):,.2f}",
                'Type': 'Forecast'
            })
    
    # Create two separate tables side-by-side: Historical (Actual) and Forecast
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### 📈 Historical (Actual) Sales")
        historical_data = [d for d in comparison_data if d.get('Type') == 'Historical']
        
        if historical_data:
            # Create clean table
            hist_table = []
            for idx, item in enumerate(historical_data, 1):
                hist_table.append({
                    'Sr.': idx,
                    'Date': item['Date'],
                    'Actual': item['Actual']
                })
            
            hist_df_display = pd.DataFrame(hist_table)
            st.dataframe(hist_df_display, use_container_width=True, hide_index=True)
        else:
            st.info("No historical data available")
    
    with col2:
        st.markdown("#### 🔮 Forecasted Sales")
        forecast_data = [d for d in comparison_data if d.get('Type') == 'Forecast']
        
        if forecast_data:
            # Create clean table
            fore_table = []
            for idx, item in enumerate(forecast_data, 1):
                fore_table.append({
                    'Sr.': idx,
                    'Date': item['Date'],
                    'Forecast': item['Forecast']
                })
            
            fore_df_display = pd.DataFrame(fore_table)
            st.dataframe(fore_df_display, use_container_width=True, hide_index=True)
        else:
            st.info("No forecast data available")
    
    # Download comparison data
    st.markdown("#### 📥 Export Comparison Data")
    
    col_dl1, col_dl2 = st.columns(2)
    
    with col_dl1:
        if historical_data or forecast_data:
            # Create combined DataFrame for download
            all_comparison = pd.DataFrame(comparison_data)
            csv = all_comparison.to_csv(index=False)
            
            st.download_button(
                label="📊 Download Full Comparison (CSV)",
                data=csv,
                file_name=f"forecast_vs_actual_{datetime.now().strftime('%Y%m%d')}.csv",
                mime="text/csv"
            )
else:
    st.warning("No forecast results available. Please run forecasting models first.")



def _combine_all_forecasts(results):
    all_forecasts = []
    for model_name, result in results.items():
        if 'forecast' in result:
            forecast_df = result['forecast'].copy()
            forecast_df['Model'] = model_name
            # Add metrics if available
            if 'metrics' in result:
                metrics = result['metrics']
                forecast_df['MAE'] = metrics.get('mae', '')
                forecast_df['RMSE'] = metrics.get('rmse', '')
                forecast_df['MAPE'] = metrics.get('mape', '')
            all_forecasts.append(forecast_df)
    return pd.concat(all_forecasts, ignore_index=True) if all_forecasts else None

# Note: PDF generation removed per request. Word (.docx) export is supported via python-docx.

def _embed_plotly_image(doc, fig, width_inches=6):
    """Helper to embed a Plotly figure into a docx Document"""
    try:
        import plotly.io as pio
        # Use kaleido for static image export
        img_bytes = pio.to_image(fig, format='png', engine='kaleido', scale=2)
        img_buf = io.BytesIO(img_bytes)
        doc.add_picture(img_buf, width=Inches(width_inches))
        return True
    except Exception as e:
        doc.add_paragraph(f"[Chart could not be generated: {str(e)}. Ensure 'kaleido' is installed.]")
        return False



def generate_docx_bytes(results, config, report_sections, facility_options=[], gis_options=[]):
    """Generate a .docx report (bytes). Embeds tables and charts (PNG) if dependencies available."""
    if not docx_available:
        st.warning("python-docx not available. Install 'python-docx' to enable Word exports.")
        return None

    doc = Document()
    doc.add_heading('Retail Sales Forecasting Report', level=1)

    # Metadata table
    combined = _combine_all_forecasts(results)
    total_records = len(combined) if combined is not None else 0
    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.style = 'LightShading-Accent1'
    meta_table.cell(0,0).text = 'Report Type:'
    meta_table.cell(0,1).text = 'Custom Report'
    meta_table.cell(1,0).text = 'Generated On:'
    meta_table.cell(1,1).text = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    meta_table.cell(2,0).text = 'Data Records:'
    meta_table.cell(2,1).text = str(total_records)

    doc.add_paragraph('')

    # Helper: prepare x values (datetime or index)
    def _prepare_x(df, xcol='ds'):
        if xcol in df.columns:
            x = pd.to_datetime(df[xcol], errors='coerce')
            if x.isna().all():
                return pd.Series(range(len(df))), False
            return x, True
        return pd.Series(range(len(df))), False

    # Initialize best_model tracking (used later even if metrics not included)
    best_model = None
    best_mae = float('inf')

    # Model Performance Summary
    if "Technical Details" in report_sections:
        doc.add_heading('Model Performance Summary', level=2)
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = 'LightShading-Accent1'
        hdr_cells = tbl.rows[0].cells
        hdr_cells[0].text = 'Model'
        hdr_cells[1].text = 'MAE'
        hdr_cells[2].text = 'RMSE'
        hdr_cells[3].text = 'MAPE'
        hdr_cells[4].text = 'R2'
        # make header text bold
        try:
            for c in hdr_cells:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
        except Exception:
            pass
        for model_name, result in results.items():
            m = result.get('metrics', {})
            row_cells = tbl.add_row().cells
            row_cells[0].text = str(model_name)
            # Format numbers: MAE/RMSE 4 decimals, MAPE 3 decimals, R2 4 decimals
            try:
                mae_val = m.get('mae', None)
                row_cells[1].text = f"{float(mae_val):.4f}" if mae_val is not None and mae_val != '' else ''
            except Exception:
                row_cells[1].text = str(m.get('mae', ''))
            try:
                rmse_val = m.get('rmse', None)
                row_cells[2].text = f"{float(rmse_val):.4f}" if rmse_val is not None and rmse_val != '' else ''
            except Exception:
                row_cells[2].text = str(m.get('rmse', ''))
            try:
                mape_val = m.get('mape', None)
                row_cells[3].text = f"{float(mape_val):.3f}%" if mape_val is not None and mape_val != '' else ''
            except Exception:
                row_cells[3].text = str(m.get('mape', ''))
            try:
                r2_val = m.get('r2', None)
                row_cells[4].text = f"{float(r2_val):.4f}" if r2_val is not None and r2_val != '' else ''
            except Exception:
                row_cells[4].text = str(m.get('r2', ''))
            try:
                if m.get('mae') is not None and float(m.get('mae')) < best_mae:
                    best_mae = float(m.get('mae'))
                    best_model = model_name
            except Exception:
                pass
    else:
        # Still need to find best model for other sections
        for model_name, result in results.items():
            m = result.get('metrics', {})
            try:
                if m.get('mae') is not None and float(m.get('mae')) < best_mae:
                    best_mae = float(m.get('mae'))
                    best_model = model_name
            except Exception:
                pass

    # Forecast Summary - Best Model
    if "Executive Summary" in report_sections and best_model and best_model in results and 'forecast' in results[best_model]:
        doc.add_heading(f'Forecast Summary - Best Model: {best_model}', level=2)
        fdf = results[best_model]['forecast'].copy()
        try:
            fdf['ds'] = pd.to_datetime(fdf['ds'])
        except Exception:
            pass
        total_forecast = fdf['yhat'].sum() if 'yhat' in fdf.columns else 0
        avg_daily = fdf['yhat'].mean() if 'yhat' in fdf.columns else 0
        peak_date = fdf.loc[fdf['yhat'].idxmax(), 'ds'] if 'yhat' in fdf.columns and not fdf['yhat'].isna().all() else None
        min_date = fdf.loc[fdf['yhat'].idxmin(), 'ds'] if 'yhat' in fdf.columns and not fdf['yhat'].isna().all() else None

        sum_tbl = doc.add_table(rows=4, cols=2)
        sum_tbl.cell(0,0).text = 'Total Forecast'
        sum_tbl.cell(0,1).text = f"₹{total_forecast:,.2f}"
        sum_tbl.cell(1,0).text = 'Average Daily'
        sum_tbl.cell(1,1).text = f"₹{avg_daily:,.2f}" if pd.notnull(avg_daily) else 'N/A'
        sum_tbl.cell(2,0).text = 'Peak Day'
        sum_tbl.cell(2,1).text = peak_date.strftime('%Y-%m-%d') if hasattr(peak_date, 'strftime') else 'N/A'
        sum_tbl.cell(3,0).text = 'Lowest Day'
        sum_tbl.cell(3,1).text = min_date.strftime('%Y-%m-%d') if hasattr(min_date, 'strftime') else 'N/A'

        # Forecast Visualizations: Sales Forecast Comparison (historical + predicted) and Most Sold Category
        if "Visualizations" in report_sections:
            doc.add_heading('Forecast Visualizations', level=2)
            try:
                import plotly.io as pio
                # Build Sales Forecast Comparison chart using user-selected columns
                fig_sales = go.Figure()
                x_is_date = False
                palette = px.colors.qualitative.Plotly if hasattr(px.colors, 'qualitative') else ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd']
                date_col = config.get('date_col', 'ds')
                val_col = config.get('value_col', 'y')
                for i, (model_name, res) in enumerate(results.items()):
                    color = palette[i % len(palette)]
                    # historical
                    if 'historical' in res:
                        h = res['historical'].copy()
                        if date_col in h.columns:
                            h[date_col] = pd.to_datetime(h[date_col], errors='coerce')
                        if val_col in h.columns:
                            h[val_col] = pd.to_numeric(h[val_col], errors='coerce')
                        if date_col in h.columns and h[date_col].notna().any():
                            h = h.sort_values(date_col)
                            fig_sales.add_trace(go.Scatter(x=h[date_col], y=h[val_col], mode='lines', name=f'{model_name} - Historical', line=dict(color=color, width=3)))
                            x_is_date = True
                        elif val_col in h.columns and h[val_col].notna().any():
                            fig_sales.add_trace(go.Scatter(x=h.index, y=h[val_col], mode='lines', name=f'{model_name} - Historical', line=dict(color=color, width=3)))

                    # forecast
                    if 'forecast' in res:
                        f = res['forecast'].copy()
                        if date_col in f.columns:
                            f[date_col] = pd.to_datetime(f[date_col], errors='coerce')
                        if 'yhat' in f.columns:
                            f['yhat'] = pd.to_numeric(f['yhat'], errors='coerce')
                        # draw CI behind lines if available
                        if date_col in f.columns and f[date_col].notna().any():
                            f = f.sort_values(date_col)
                            if 'yhat_lower' in f.columns and 'yhat_upper' in f.columns:
                                try:
                                    low = pd.to_numeric(f['yhat_lower'], errors='coerce')
                                    up = pd.to_numeric(f['yhat_upper'], errors='coerce')
                                    if low.notna().any() and up.notna().any():
                                        xvals = f[date_col]
                                        fig_sales.add_trace(go.Scatter(x=list(xvals) + list(xvals)[::-1], y=list(up) + list(low)[::-1], fill='toself', fillcolor='rgba(100,100,150,0.08)', line=dict(color='rgba(255,255,255,0)'), name=f'{model_name} - CI', showlegend=False))
                                except Exception:
                                    pass
                            fig_sales.add_trace(go.Scatter(x=f[date_col], y=f['yhat'], mode='lines', name=f'{model_name} - Forecast', line=dict(color=color, dash='dash', width=3)))
                            x_is_date = True
                        elif 'yhat' in f.columns and f['yhat'].notna().any():
                            fig_sales.add_trace(go.Scatter(x=f.index, y=f['yhat'], mode='lines', name=f'{model_name} - Forecast', line=dict(color=color, dash='dash', width=3)))

                if fig_sales.data:
                    if x_is_date:
                        fig_sales.update_layout(xaxis=dict(title='Date', type='date'))
                    else:
                        fig_sales.update_layout(xaxis=dict(title='Index'))
                    fig_sales.update_layout(title='Sales Forecast Comparison', yaxis=dict(title='Sales', tickformat=',.2f'), height=480)
                else:
                    fig_sales = None

                # Most sold category of product (if available in combined forecasts or original data)
                combined = _combine_all_forecasts(results)
                cat_df = None
                if combined is not None and 'Category' in combined.columns and 'yhat' in combined.columns:
                    cat_df = combined.groupby('Category', as_index=False)['yhat'].sum().sort_values('yhat', ascending=False)
                else:
                    original = st.session_state.get('uploaded_data')
                    if original is not None and 'Category' in original.columns and ('Sales' in original.columns or 'y' in original.columns):
                        sales_col = 'Sales' if 'Sales' in original.columns else 'y'
                        cat_df = original.groupby('Category', as_index=False)[sales_col].sum().sort_values(sales_col, ascending=False)

                if cat_df is not None and not cat_df.empty:
                    # standardize column names for plotting
                    if 'yhat' in cat_df.columns:
                        val_col_cat = 'yhat'
                    elif 'Sales' in cat_df.columns:
                        val_col_cat = 'Sales'
                    else:
                        val_col_cat = cat_df.columns[1]

                    fig_cat = px.bar(cat_df, x='Category', y=val_col_cat, title='Most Sold Category', color='Category')
                    fig_cat.update_layout(height=350, showlegend=False)
                else:
                    fig_cat = None

                # Embed two charts side-by-side using a 1x2 table
                try:
                    tbl = doc.add_table(rows=1, cols=2)
                    tbl.autofit = True
                    cell0 = tbl.cell(0,0)
                    cell1 = tbl.cell(0,1)
                    if fig_sales is not None:
                        img_bytes = pio.to_image(fig_sales, format='png', engine='kaleido')
                        img_buf = io.BytesIO(img_bytes)
                        para = cell0.paragraphs[0]
                        run = para.add_run()
                        run.add_picture(img_buf, width=Inches(3.8))
                    else:
                        cell0.text = 'No Sales chart available.'
                    if fig_cat is not None:
                        img_bytes = pio.to_image(fig_cat, format='png', engine='kaleido')
                        img_buf = io.BytesIO(img_bytes)
                        para2 = cell1.paragraphs[0]
                        run2 = para2.add_run()
                        run2.add_picture(img_buf, width=Inches(3.8))
                    else:
                        cell1.text = 'No Category chart available.'
                except Exception:
                    # fallback to sequential images
                    if fig_sales is not None:
                        img_bytes = pio.to_image(fig_sales, format='png', engine='kaleido')
                        img_buf = io.BytesIO(img_bytes)
                        doc.add_picture(img_buf, width=Inches(6))
                    if fig_cat is not None:
                        img_bytes = pio.to_image(fig_cat, format='png', engine='kaleido')
                        img_buf = io.BytesIO(img_bytes)
                        doc.add_picture(img_buf, width=Inches(6))
            except Exception:
                doc.add_paragraph('Charts not embedded: ensure kaleido and python-docx are installed.')

    #Recommendations Section
    if "Recommendations" in report_sections:
        doc.add_heading('AI-Powered Business Recommendations', level=2)

        # Get data for AI analysis
        data_df = st.session_state['uploaded_data'] if st.session_state.get('uploaded_data') is not None else st.session_state.get('processed_data')
        if data_df is not None:
            insights = analyze_business_metrics(data_df, results)

            # Prepare forecast data for AI
            forecast_data = {
                'total_forecast': insights.get('total_forecast', 0),
                'avg_daily': insights.get('avg_forecast', 0),
                'best_model': insights.get('best_model', 'N/A'),
                'period': config.get('forecast_days', 30),
                'volatility': 'High' if insights.get('volatility', 0) > 0.3 else 'Moderate' if insights.get('volatility', 0) > 0.15 else 'Low',
                'growth_rate': insights.get('growth_rate', 0)
            }

            # Business context
            business_context = f"""
            Business Type: Retail Store
            Data Period: {len(data_df)} transactions
            Categories: {data_df['Category'].nunique() if 'Category' in data_df.columns else 'N/A'} product categories
            Stores: {data_df['Store'].nunique() if 'Store' in data_df.columns else 'Single store'}
            Average Transaction Value: ₹{insights.get('avg_daily', 0):,.0f}
            Sales Volatility: {forecast_data['volatility']}
            Forecasted Growth: {forecast_data['growth_rate']:.1f}%
            """

            # Get AI recommendations (Gemini AI-powered)
            ai_recommendations = get_ai_recommendations(forecast_data, business_context)

            doc.add_paragraph('Based on analysis of your sales data and forecasts, here are personalized recommendations:')
            doc.add_paragraph('')

            for i, rec in enumerate(ai_recommendations[:7], 1):
                # Clean up recommendation text (remove excessive icons)
                clean_rec = rec.replace('📦', '').replace('👥', '').replace('💰', '').replace('📈', '').replace('🎯', '').replace('📊', '').replace('🔄', '')
                clean_rec = clean_rec.replace('**', '')  # Remove bold formatting
                doc.add_paragraph(f'{i}. {clean_rec.strip()}')



        # Business Insights Section
        if "Business Insights" in report_sections:
            doc.add_heading('Business Intelligence Insights', level=2)

            # Key insights table
            insights_table = doc.add_table(rows=5, cols=2)
            insights_table.style = 'LightShading-Accent1'

            insights_data = [
                ('Total Historical Sales', f"₹{insights.get('total_sales', 0):,.0f}"),
                ('Average Daily Sales', f"₹{insights.get('avg_daily', 0):,.0f}"),
                ('Sales Volatility', f"{insights.get('volatility', 0):.1%}"),
                ('Forecasted Growth Rate', f"{insights.get('growth_rate', 0):.1f}%"),
                ('Best Performing Model', insights.get('best_model', 'N/A'))
            ]

            for i, (metric, value) in enumerate(insights_data):
                insights_table.cell(i, 0).text = metric
                insights_table.cell(i, 1).text = str(value)

        # Category Performance (if available)
        if "Category Analysis" in report_sections:
            if 'Category' in data_df.columns:
                doc.add_heading('Category Performance Analysis', level=3)
                category_metrics = data_df.groupby('Category')['Sales'].agg(['sum', 'mean', 'count']).round(2)
                category_metrics.columns = ['Total Sales', 'Avg Sales', 'Transactions']

                # Top 5 categories table
                top_5 = category_metrics.nlargest(5, 'Total Sales')
                cat_table = doc.add_table(rows=1 + len(top_5), cols=4)
                cat_table.style = 'LightShading-Accent1'

                # Headers
                headers = ['Category', 'Total Sales', 'Avg Sales', 'Transactions']
                for j, header in enumerate(headers):
                    cat_table.cell(0, j).text = header

                # Data rows
                for i, (cat, row) in enumerate(top_5.iterrows(), 1):
                    cat_table.cell(i, 0).text = str(cat)
                    cat_table.cell(i, 1).text = f"₹{row['Total Sales']:,.0f}"
                    cat_table.cell(i, 2).text = f"₹{row['Avg Sales']:,.0f}"
                    cat_table.cell(i, 3).text = f"{row['Transactions']:.0f}"

            # Store Performance (if available)
            if 'Store' in data_df.columns:
                doc.add_heading('Store Performance Analysis', level=3)
                store_metrics = data_df.groupby('Store')['Sales'].agg(['sum', 'mean', 'count']).round(2)
                store_metrics.columns = ['Total Sales', 'Avg Sales', 'Transactions']

                store_table = doc.add_table(rows=1 + len(store_metrics), cols=4)
                store_table.style = 'LightShading-Accent1'

                # Headers
                headers = ['Store', 'Total Sales', 'Avg Sales', 'Transactions']
                for j, header in enumerate(headers):
                    store_table.cell(0, j).text = header

                # Data rows
                for i, (store, row) in enumerate(store_metrics.iterrows(), 1):
                    store_table.cell(i, 0).text = str(store)
                    store_table.cell(i, 1).text = f"₹{row['Total Sales']:,.0f}"
                    store_table.cell(i, 2).text = f"₹{row['Avg Sales']:,.0f}"
                    store_table.cell(i, 3).text = f"{row['Transactions']:.0f}"


    # Facility Layout Section
    if facility_options:
        doc.add_heading('Facility Layout Analysis', level=2)
        
        if "Optimized Layout" in facility_options and 'departments' in st.session_state:
            doc.add_heading('Department Configuration', level=3)
            depts = st.session_state.departments
            dept_table = doc.add_table(rows=1 + len(depts), cols=3)
            dept_table.style = 'LightShading-Accent1'
            
            headers = ['Department', 'Area (sq ft)', 'Type']
            for j, h in enumerate(headers):
                dept_table.cell(0, j).text = h
                
            for i, (_, row) in enumerate(depts.iterrows(), 1):
                dept_table.cell(i, 0).text = str(row['Name'])
                dept_table.cell(i, 1).text = f"{row['Area']:,.0f}"
                dept_table.cell(i, 2).text = str(row['Type'])
            doc.add_paragraph('')

            # Diagram
            if 'layout_positions' in st.session_state:
                doc.add_heading('Optimized Layout Diagram', level=4)
                try:
                    positions = st.session_state.layout_positions
                    fig = go.Figure()
                    
                    # Calculate facility dimensions dynamically
                    max_x = positions['X'].max() + positions['Width'].max() / 2
                    max_y = positions['Y'].max() + positions['Height'].max() / 2
                    facility_width = max_x * 1.2  # Add 20% margin
                    facility_length = max_y * 1.2
                    
                    # Draw facility boundary
                    fig.add_shape(
                        type="rect",
                        x0=0, y0=0,
                        x1=facility_width, y1=facility_length,
                        line=dict(color="black", width=3)
                    )
                    
                    # Draw departments
                    for _, dept in positions.iterrows():
                        # Rectangle for department
                        fig.add_shape(
                            type="rect", 
                            x0=dept['X'] - dept['Width']/2, 
                            y0=dept['Y'] - dept['Height']/2,
                            x1=dept['X'] + dept['Width']/2, 
                            y1=dept['Y'] + dept['Height']/2,
                            fillcolor='lightblue', 
                            line=dict(color='darkblue', width=2), 
                            opacity=0.7
                        )
                        
                        # Label
                        fig.add_annotation(
                            x=dept['X'], 
                            y=dept['Y'], 
                            text=dept['Department'], 
                            showarrow=False, 
                            font=dict(size=12, color='black', family='Arial Black')
                        )
                    
                    # Draw relationship connections if available
                    if 'relationship_matrix' in st.session_state:
                        closeness_colors = {
                            'A': 'darkgreen',
                            'E': 'green',
                            'I': 'lightgreen',
                            'O': 'gray',
                            'U': 'lightgray',
                            'X': 'red'
                        }
                        
                        for _, rel in st.session_state.relationship_matrix.iterrows():
                            from_pos = positions[positions['Department'] == rel['From']]
                            to_pos = positions[positions['Department'] == rel['To']]
                            
                            if len(from_pos) > 0 and len(to_pos) > 0 and rel['Closeness'] in ['A', 'E', 'I', 'X']:
                                fig.add_trace(go.Scatter(
                                    x=[from_pos.iloc[0]['X'], to_pos.iloc[0]['X']],
                                    y=[from_pos.iloc[0]['Y'], to_pos.iloc[0]['Y']],
                                    mode='lines',
                                    line=dict(
                                        color=closeness_colors.get(rel['Closeness'], 'gray'),
                                        width=2,
                                        dash='dot'
                                    ),
                                    showlegend=False,
                                    hoverinfo='none'
                                ))
                        
                        # Add legend for relationship lines - positioned on the right side
                        legend_x = facility_width * 0.68  # Right side instead of left
                        legend_y = facility_length * 0.98  # Top
                        legend_items = [
                            ('A = Absolutely Necessary', 'darkgreen'),
                            ('E = Especially Important', 'green'),
                            ('I = Important', 'lightgreen'),
                            ('O = Optional', 'gray'),
                            ('U = Undesirable', 'lightgray'),
                            ('X = Undesirable', 'red')
                        ]
                        
                        # Add legend box background
                        fig.add_shape(
                            type="rect",
                            x0=legend_x - 1, y0=legend_y - len(legend_items) * 5 - 3,
                            x1=legend_x + 30, y1=legend_y + 1,
                            fillcolor='white',
                            line=dict(color='black', width=1.5),
                            opacity=0.95
                        )
                        
                        # Add legend title
                        fig.add_annotation(
                            x=legend_x + 0.5, y=legend_y - 1,
                            text="<b>Relationship Lines:</b>",
                            showarrow=False,
                            font=dict(size=8, color='black', family='roman'),
                            xanchor='left',
                            yanchor='top'
                        )
                        
                        # Add legend items
                        for i, (label, color) in enumerate(legend_items):
                            y_pos = legend_y - (i + 1.5) * 5
                            # Color indicator line
                            fig.add_shape(
                                type="line",
                                x0=legend_x + 0.5, y0=y_pos,
                                x1=legend_x + 4, y1=y_pos,
                                line=dict(color=color, width=2.5, dash='dot')
                            )
                            # Label text
                            fig.add_annotation(
                                x=legend_x + 5, y=y_pos,
                                text=label,
                                showarrow=False,
                                font=dict(size=8, color='black', family='roman'),
                                xanchor='left',
                                yanchor='middle'
                            )
                    
                    fig.update_layout(
                        title="Facility Layout with Department Relationships",
                        xaxis=dict(title="Width (ft)", range=[0, facility_width], showgrid=True, gridcolor='lightgray'),
                        yaxis=dict(title="Length (ft)", range=[0, facility_length], showgrid=True, gridcolor='lightgray'),
                        height=600,
                        width=800,
                        showlegend=False
                    )
                    _embed_plotly_image(doc, fig, width_inches=6.5)
                except Exception as e:
                    doc.add_paragraph(f"[Layout diagram error: {str(e)}]")
            doc.add_paragraph('')

        if "Relationship Network" in facility_options and 'relationship_matrix' in st.session_state:
            doc.add_heading('Activity Relationships', level=3)
            rels = st.session_state.relationship_matrix
            # Filter for important relationships only (A, E, I, X)
            important_rels = rels[rels['Closeness'].isin(['A', 'E', 'I', 'X'])]
            
            if not important_rels.empty:
                rel_table = doc.add_table(rows=1 + len(important_rels), cols=3)
                rel_table.style = 'LightShading-Accent1'
                
                headers = ['From', 'To', 'Closeness']
                for j, h in enumerate(headers):
                    rel_table.cell(0, j).text = h
                    
                for i, (_, row) in enumerate(important_rels.iterrows(), 1):
                    rel_table.cell(i, 0).text = str(row['From'])
                    rel_table.cell(i, 1).text = str(row['To'])
                    rel_table.cell(i, 2).text = str(row['Closeness'])
                doc.add_paragraph('Note: Showing only Critical (A, E), Important (I), and Undesirable (X) relationships.')
            else:
                doc.add_paragraph('No critical or important relationships defined.')
            doc.add_paragraph('')

            # Diagram
            doc.add_heading('Relationship Network Diagram', level=4)
            try:
                # Create network graph with better styling
                G = nx.Graph()
                dept_names = list(set(rels['From'].unique()) | set(rels['To'].unique()))
                for d in dept_names: 
                    G.add_node(d)
                
                # Define relationship colors and weights
                closeness_colors_map = {
                    'A': ('darkgreen', 4),
                    'E': ('green', 3),
                    'I': ('lightgreen', 2),
                    'O': ('gray', 1),
                    'U': ('lightgray', 1),
                    'X': ('red', 3)
                }
                
                # Add edges with colors
                edge_traces = []
                for _, row in rels.iterrows():
                    if row['Closeness'] in closeness_colors_map:
                        color, width = closeness_colors_map[row['Closeness']]
                        if row['From'] in G.nodes() and row['To'] in G.nodes():
                            G.add_edge(row['From'], row['To'])
                
                # Use spring layout with better spacing
                pos = nx.spring_layout(G, k=2, iterations=50, seed=42)
                
                # Create edge traces with colors
                for _, row in rels.iterrows():
                    if row['From'] in pos and row['To'] in pos and row['Closeness'] in closeness_colors_map:
                        color, width = closeness_colors_map[row['Closeness']]
                        x0, y0 = pos[row['From']]
                        x1, y1 = pos[row['To']]
                        edge_traces.append(go.Scatter(
                            x=[x0, x1, None],
                            y=[y0, y1, None],
                            mode='lines',
                            line=dict(width=width, color=color),
                            showlegend=False,
                            hoverinfo='text',
                            text=f"{row['From']} - {row['To']}: {row['Closeness']}"
                        ))
                
                # Create node trace
                node_x = [pos[node][0] for node in G.nodes()]
                node_y = [pos[node][1] for node in G.nodes()]
                node_text = list(G.nodes())
                
                node_trace = go.Scatter(
                    x=node_x, 
                    y=node_y, 
                    mode='markers+text', 
                    text=node_text,
                    textposition="top center",
                    textfont=dict(size=12, color='black', family='Arial Black'),
                    marker=dict(
                        size=40,
                        color='lightblue',
                        line=dict(width=2, color='darkblue')
                    ),
                    showlegend=False,
                    hoverinfo='text',
                    hovertext=node_text
                )
                
                # Create figure
                fig = go.Figure(data=edge_traces + [node_trace])
                fig.update_layout(
                    title="Department Relationship Network",
                    showlegend=False,
                    xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
                    yaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
                    height=600,
                    width=800,
                    hovermode='closest'
                )
                _embed_plotly_image(doc, fig, width_inches=6.5)
            except Exception as e:
                doc.add_paragraph(f"[Network diagram error: {str(e)}]")
            doc.add_paragraph('')

        if "Material Flow" in facility_options and 'flow_matrix' in st.session_state:
            doc.add_heading('Material Flow Analysis', level=3)
            flows = st.session_state.flow_matrix
            active_flows = flows[flows['Flow Volume'] > 0].sort_values('Flow Volume', ascending=False)
            
            if not active_flows.empty:
                flow_table = doc.add_table(rows=1 + len(active_flows), cols=3)
                flow_table.style = 'LightShading-Accent1'
                
                headers = ['From', 'To', 'Volume']
                for j, h in enumerate(headers):
                    flow_table.cell(0, j).text = h
                    
                for i, (_, row) in enumerate(active_flows.iterrows(), 1):
                    flow_table.cell(i, 0).text = str(row['From'])
                    flow_table.cell(i, 1).text = str(row['To'])
                    flow_table.cell(i, 2).text = f"{row['Flow Volume']} {row['Unit']}"
            else:
                doc.add_paragraph('No material flow volumes defined.')
            doc.add_paragraph('')

            # Diagram
            doc.add_heading('Material Flow Diagram', level=4)
            try:
                active = flows[flows['Flow Volume'] > 0]
                if not active.empty:
                    nodes = list(set(active['From'].unique()) | set(active['To'].unique()))
                    node_map = {n:i for i,n in enumerate(nodes)}
                    
                    fig = go.Figure(data=[go.Sankey(
                        node=dict(label=nodes, pad=15, thickness=20, color="lightblue"),
                        link=dict(
                            source=[node_map[x] for x in active['From']],
                            target=[node_map[x] for x in active['To']],
                            value=active['Flow Volume']
                        ))])
                    fig.update_layout(title="Material Flow Volume", font_size=10)
                    _embed_plotly_image(doc, fig)
            except Exception as e:
                doc.add_paragraph(f"[Flow diagram error: {str(e)}]")
            doc.add_paragraph('')

    # Store Location GIS Section
    if gis_options:
        doc.add_heading('Store Location GIS Analysis', level=2)
        
        if "Network Statistics" in gis_options:
            doc.add_heading('Network Statistics', level=3)
            stats_table = doc.add_table(rows=4, cols=2)
            stats_table.style = 'LightShading-Accent1'
            
            reliance_count = len(st.session_state.reliance_stores) if 'reliance_stores' in st.session_state else 0
            competitor_count = len(st.session_state.competitor_stores) if 'competitor_stores' in st.session_state else 0
            
            stats_data = [
                ('Reliance Fresh (Competitor)', str(reliance_count)),
                ('KPN Fresh (Target)', str(competitor_count))
            ]
            
            if 'population_data' in st.session_state:
                pop_df = st.session_state.population_data
                if 'Population' in pop_df.columns:
                    total_pop = pop_df['Population'].sum()
                    avg_pop = pop_df['Population'].mean()
                    stats_data.append(('Total Population Coverage', f"{total_pop:,.0f}"))
                    stats_data.append(('Average Ward Population', f"{avg_pop:,.0f}"))
            
            for i, (metric, value) in enumerate(stats_data):
                if i < 4: # Ensure we don't exceed table rows
                    stats_table.cell(i, 0).text = metric
                    stats_table.cell(i, 1).text = value
            doc.add_paragraph('')

        if "Store Locations" in gis_options and 'reliance_stores' in st.session_state:
            # KPN Fresh Stores (Target)
            if 'competitor_stores' in st.session_state:
                doc.add_heading('Current KPN Fresh Locations', level=3)
                kpn_stores = st.session_state.competitor_stores
                display_kpn = kpn_stores.head(20)
                
                kpn_table = doc.add_table(rows=1 + len(display_kpn), cols=2)
                kpn_table.style = 'LightShading-Accent1'
                
                headers = ['Store Name', 'Address']
                for j, h in enumerate(headers):
                    kpn_table.cell(0, j).text = h
                    
                for i, (_, row) in enumerate(display_kpn.iterrows(), 1):
                    kpn_table.cell(i, 0).text = str(row.get('name', 'N/A'))
                    kpn_table.cell(i, 1).text = str(row.get('address', 'N/A'))[:50] + '...'
                
                if len(kpn_stores) > 20:
                    doc.add_paragraph(f'...and {len(kpn_stores) - 20} more KPN stores.')
                doc.add_paragraph('')



            # Map
            doc.add_heading('Store Location Map', level=4)
            try:
                fig = go.Figure()
                
                # Reliance Stores
                stores = st.session_state.reliance_stores
                if not stores.empty:
                    # Extract lat/lon from geometry if needed, or assume columns
                    lats = [g.y for g in stores.geometry]
                    lons = [g.x for g in stores.geometry]
                    names = stores['name'].tolist()
                    fig.add_trace(go.Scattermapbox(
                        lat=lats, lon=lons, mode='markers', marker=dict(size=10, color='blue'),
                        text=names, name='Reliance Fresh'
                    ))
                
                # Competitor Stores
                if 'competitor_stores' in st.session_state:
                    comp = st.session_state.competitor_stores
                    if not comp.empty:
                        lats = [g.y for g in comp.geometry]
                        lons = [g.x for g in comp.geometry]
                        names = comp['name'].tolist()
                        fig.add_trace(go.Scattermapbox(
                            lat=lats, lon=lons, mode='markers', marker=dict(size=8, color='red'),
                            text=names, name='KPN Fresh (Target)'
                        ))
                
                # Recommendations - New Recommended Store Locations
                if 'gis_recommendations' in st.session_state:
                    recs = st.session_state.gis_recommendations
                    if recs:
                        lats = [r['lat'] for r in recs]
                        lons = [r['lng'] for r in recs]
                        names = [f"New Recommended #{i+1}: {r['ward_name']}" for i,r in enumerate(recs)]
                        fig.add_trace(go.Scattermapbox(
                            lat=lats, lon=lons, mode='markers', 
                            marker=dict(size=14, color='limegreen', symbol='circle'),
                            text=names, name='Recommended New Stores'
                        ))

                # Center map
                center_lat = 12.9716
                center_lon = 77.5946
                if not stores.empty:
                    center_lat = np.mean([g.y for g in stores.geometry])
                    center_lon = np.mean([g.x for g in stores.geometry])

                fig.update_layout(
                    mapbox_style="open-street-map",
                    mapbox=dict(center=dict(lat=center_lat, lon=center_lon), zoom=10),
                    margin={"r":0,"t":0,"l":0,"b":0},
                    height=500,
                    showlegend=True
                )
                _embed_plotly_image(doc, fig)
            except Exception as e:
                doc.add_paragraph(f"[Map generation error: {str(e)}]")
            doc.add_paragraph('')

        if "Strategic Recommendations" in gis_options and 'gis_recommendations' in st.session_state:
            doc.add_heading('Strategic Location Recommendations', level=3)
            recs = st.session_state.gis_recommendations
            
            if recs:
                rec_table = doc.add_table(rows=1 + min(len(recs), 10), cols=4)
                rec_table.style = 'LightShading-Accent1'
                
                headers = ['Rank', 'Ward', 'Population', 'Score']
                for j, h in enumerate(headers):
                    rec_table.cell(0, j).text = h
                    
                for i, rec in enumerate(recs[:10], 1):
                    rec_table.cell(i, 0).text = str(i)
                    rec_table.cell(i, 1).text = f"{rec['ward_name']} ({rec['assembly']})"
                    rec_table.cell(i, 2).text = f"{rec['population']:,}"
                    rec_table.cell(i, 3).text = f"{rec['total_score']:.1f}"
                doc.add_paragraph('Top 10 AI-generated location recommendations based on population, competition, and accessibility.')
            else:
                doc.add_paragraph('No recommendations generated yet.')
            doc.add_paragraph('')

    # Save to bytes
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

with col1:
    if st.button("📊 Generate Excel Report", width='stretch'):
        # Create Excel report
        output = io.BytesIO()
        
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            # Summary sheet
            summary_data = {
                'Metric': ['Report Generated', 'Best Model', 'Forecast Period', 'Total Models'],
                'Value': [
                    datetime.now().strftime('%Y-%m-%d %H:%M'),
                    best_model if 'best_model' in locals() else 'N/A',
                    f"{config.get('forecast_days', 'N/A')} days",
                    len(results)
                ]
            }
            summary_df = pd.DataFrame(summary_data)
            summary_df.to_excel(writer, sheet_name='Summary', index=False)
            
            # Metrics sheet
            if 'metrics_data' in locals() and metrics_data:
                metrics_df.to_excel(writer, sheet_name='Model_Metrics', index=False)
            
            # Forecasts sheet
            all_forecasts = []
            for model_name, result in results.items():
                if 'forecast' in result:
                    forecast_df = result['forecast'].copy()
                    forecast_df['Model'] = model_name
                    all_forecasts.append(forecast_df)
            
            if all_forecasts:
                combined_forecasts = pd.concat(all_forecasts, ignore_index=True)
                combined_forecasts.to_excel(writer, sheet_name='Forecasts', index=False)
        
        output.seek(0)
        
        st.download_button(
            label="📥 Download Excel Report",
            data=output.getvalue(),
            file_name=f"retail_forecast_report_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

with col2:
    if st.button("📄 Generate CSV Export", width='stretch'):
        combined_df = _combine_all_forecasts(results)
        if combined_df is not None:
            csv = combined_df.to_csv(index=False)
            st.download_button(
                label="📥 Download CSV Data",
                data=csv,
                file_name=f"forecast_data_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                mime="text/csv"
            )

    pass

# Document generation button (Word)
st.markdown("---")
doc_col1, doc_col2 = st.columns([1,2])
with doc_col1:
    if st.button("📥 Download Document (.docx)"):
        docx_bytes = generate_docx_bytes(
            results, config, report_sections, 
            facility_options=facility_options if include_facility else [],
            gis_options=gis_options if include_gis else []
        )
        if docx_bytes:
            st.download_button(
                label="📥 Download Word Report",
                data=docx_bytes,
                file_name=f"retail_forecast_report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        else:
            st.error("Failed to generate Word document. Ensure 'python-docx' and 'kaleido' are installed in your environment.")


